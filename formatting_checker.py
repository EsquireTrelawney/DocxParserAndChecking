from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import difflib

from formatting_utils import (
    get_effective_first_line_indent_obj,
    get_effective_alignment,
    get_first_line_indent_cm,
    get_run_font_name,
    get_run_font_size_pt,
    get_run_font_color_rgb,
    get_run_bold_status
)

# --- Константы ---
STRUCTURAL_HEADINGS_KEYWORDS = [
    "СОДЕРЖАНИЕ", "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", 
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 
    "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", "СПИСОК ЛИТЕРАТУРЫ"
]
HEADING_1_STYLE_NAMES = ["heading 1", "заголовок 1", "header 1", "title 1"]
HEADING_2_STYLE_NAMES = ["heading 2", "заголовок 2", "header 2", "title 2"]

# --- Утилиты определения типа элемента ---

def get_paragraph_style_name(para):
    """Решил сделать функцию для безопасного получения имени стиля абзаца."""
    if para.style and para.style.name:
        return para.style.name.lower()
    return ""

def check_all_runs_are_bold(para):
    """Проверяет, что все непустые runs в абзаце эффективно полужирные."""
    if not para.runs and para.text.strip():
        return False 
    
    text_runs_exist = False
    for run in para.runs:
        if run.text.strip():
            text_runs_exist = True
            # Получаем эффективный статус жирности, учитывая стиль абзаца
            is_bold = get_run_bold_status(run, para.style) 
            if is_bold is None: is_bold = False
            if not is_bold:
                return False 
    return text_runs_exist

def is_structural_heading_type(para_text_upper_stripped):
    """Определяет, является ли текст структурным заголовком (ВВЕДЕНИЕ и т.д.)."""
    return para_text_upper_stripped in STRUCTURAL_HEADINGS_KEYWORDS

def is_appendix_heading_type(para_text_upper_stripped):
    """Определяет, является ли текст заголовком приложения ('ПРИЛОЖЕНИЕ А')."""
    return bool(re.fullmatch(r"ПРИЛОЖЕНИЕ\s+[А-ЯЁ]{1,2}", para_text_upper_stripped))

def is_in_table(para, doc):
    """Моя функция для проверки, находится ли параграф в таблице."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if para in cell.paragraphs:
                    return True
    return False

def is_main_heading(para):
    """Решил сделать такую проверку для заголовков основных разделов."""
    # Проверяем стиль параграфа на соответствие заголовку
    if hasattr(para, 'style') and para.style:
        style_name = para.style.name.lower() if hasattr(para.style, 'name') else ""
        heading_style_indicators = ["heading", "header", "title", "заголовок", "оглавление"]
        
        # Если стиль содержит признаки заголовка и при этом не имеет числовой части
        is_heading_style = any(indicator in style_name for indicator in heading_style_indicators)
        has_numeric_level = any(str(i) in style_name for i in range(1, 10))
        
        if is_heading_style and not has_numeric_level:
            # Проверяем, есть ли среди основных заголовков соответствие тексту параграфа
            cleaned_text = para.text.strip().upper()
            if cleaned_text.endswith('.'):
                cleaned_text = cleaned_text[:-1]
                
            if cleaned_text in STRUCTURAL_HEADINGS_KEYWORDS:
                return True
            
            # Проверка на частичное совпадение
            for heading in STRUCTURAL_HEADINGS_KEYWORDS:
                if heading in cleaned_text:
                    return True
    
    # Если не определили по стилю, проверяем по содержанию и форматированию
    cleaned_text = para.text.strip().upper()
    if cleaned_text.endswith('.'):
        cleaned_text = cleaned_text[:-1]
    
    # Точное совпадение с известными заголовками
    if cleaned_text in STRUCTURAL_HEADINGS_KEYWORDS:
        # Дополнительно проверяем признаки заголовка - центрирование и жирный шрифт
        is_centered = False
        is_bold = False
        
        # Проверка выравнивания
        alignment = get_effective_alignment(para)
        is_centered = (alignment == WD_ALIGN_PARAGRAPH.CENTER)
        
        # Проверка жирного шрифта
        is_bold = check_all_runs_are_bold(para)
        
        # Если соответствует тексту и имеет признаки форматирования заголовка
        if is_centered or is_bold:
            return True
        
        # Если ни один из признаков форматирования не обнаружен, 
        # но текст точно соответствует заголовку - вероятно, это заголовок
        return True
    
    return False

def is_introduction_heading(para):
    """Специальная проверка для заголовка ВВЕДЕНИЕ (нужна для активации проверок)"""
    cleaned_text = para.text.strip().upper()
    if cleaned_text.endswith('.'):
        cleaned_text = cleaned_text[:-1]
    return cleaned_text == "ВВЕДЕНИЕ"

def is_bibliography_heading(para):
    """
    Я написал эту функцию для поиска заголовков библиографии.
    
    Распознает различные варианты названий списка литературы,
    такие как "СПИСОК ЛИТЕРАТУРЫ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ".
    """
    bibliography_headings = [
        "СПИСОК ЛИТЕРАТУРЫ", 
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ",
        "БИБЛИОГРАФИЧЕСКИЙ СПИСОК",
        "БИБЛИОГРАФИЯ",
        "СПИСОК ИСТОЧНИКОВ",
        "ЛИТЕРАТУРА"
    ]
    
    # Приводим текст к верхнему регистру и убираем пробелы по краям
    text_raw = para.text.strip()
    cleaned_text = text_raw.upper()
    
    # Убираем точку в конце, если она есть
    if cleaned_text.endswith('.'):
        cleaned_text = cleaned_text[:-1]
    
    # Проверяем текст на соответствие заголовку библиографии
    is_heading = False
    for heading in bibliography_headings:
        # Полное соответствие
        if cleaned_text == heading:
            is_heading = True
            break
        # Приблизительное соответствие (например, с небольшими вариациями)
        similarity = difflib.SequenceMatcher(None, cleaned_text, heading).ratio()
        if similarity > 0.8:  # 80% совпадение
            is_heading = True
            break
    
    # Дополнительные проверки для подтверждения, что это действительно заголовок
    if is_heading:
        # Заголовки обычно короткие
        if len(text_raw) > 50:
            return False
            
        # Заголовки часто имеют отличное форматирование
        # 1. Проверка на жирный шрифт
        if check_all_runs_are_bold(para):
            return True
            
        # 2. Проверка на заглавные буквы
        if text_raw.isupper():
            return True
            
        # 3. Проверка на стиль заголовка
        if hasattr(para, 'style') and para.style:
            style_name = para.style.name.lower() if hasattr(para.style, 'name') else ""
            if 'heading' in style_name or 'заголовок' in style_name:
                return True
            
        # 4. Проверка на выравнивание по центру
        if hasattr(para, 'paragraph_format') and para.paragraph_format:
            if hasattr(para.paragraph_format, 'alignment'):
                alignment = para.paragraph_format.alignment
                if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return True
                    
    # Если текст не похож на заголовок библиографии или не имеет форматирования заголовка
    return False

def is_section_heading(para):
    """
    Мне нужно было проверять заголовки разделов вида "1. Заголовок".
    """
    # Проверка по стилю
    style_name = get_paragraph_style_name(para)
    if any(h_style in style_name for h_style in HEADING_1_STYLE_NAMES):
        return True
    
    # Проверка формата с точкой после номера (1. Заголовок)
    if re.fullmatch(r"\d{1,2}\.\s+.+", para.text.strip()):
        if check_all_runs_are_bold(para):
            return True
            
    # Проверка формата без точки после номера (1 Заголовок)
    if re.fullmatch(r"\d{1,2}\s+.+", para.text.strip()):
        if check_all_runs_are_bold(para):
            return True
            
    return False

def is_subsection_heading(para):
    """
    Проверяет, является ли параграф подзаголовком (1.1 Заголовок без точки после номера).
    
    Подзаголовок:
    1. Имеет формат "N.M Текст" или "N.M. Текст", где N и M - числа
    2. Может иметь стиль Heading 2 или аналогичный
    3. Обычно выделен жирным шрифтом
    
    Args:
        para: Объект параграфа
    
    Returns:
        bool: True если параграф является подзаголовком
    """
    # Проверка по стилю
    style_name = get_paragraph_style_name(para)
    if any(h_style in style_name for h_style in HEADING_2_STYLE_NAMES):
        return True
        
    # Проверка формата без точки после номера (правильный)
    if re.fullmatch(r"\d+(\.\d+)+\s+.+", para.text.strip()): 
         if check_all_runs_are_bold(para):
            return True
    
    # Проверка формата с точкой после номера (неправильный)
    if re.fullmatch(r"\d+(\.\d+)+\.\s+.+", para.text.strip()): 
         if check_all_runs_are_bold(para):
            return True
    
    return False

def is_figure_caption(para):
    """
    Проверяет, является ли параграф подписью к рисунку.
    
    Подпись к рисунку:
    1. Имеет формат "Рисунок N - Текст" или "Рисунок N – Текст" (с дефисом или тире)
    2. Может иметь стиль Caption или аналогичный
    3. Обычно выровнен по центру
    
    Args:
        para: Объект параграфа
    
    Returns:
        bool: True если параграф является подписью к рисунку
    """
    # Проверка по стилю
    if hasattr(para, 'style') and para.style:
        style_name = para.style.name.lower() if hasattr(para.style, 'name') else ""
        
        # Проверка на стили подписей
        caption_indicators = ["caption", "подпись", "figure", "рисунок"]
        if any(indicator in style_name for indicator in caption_indicators):
            # Если это стиль подписи и текст содержит слово "рисунок" - это подпись к рисунку
            if "рисунок" in para.text.lower() or "figure" in para.text.lower() or "рис" in para.text.lower():
                return True
    
    # Проверка по формату текста - расширенные шаблоны
    text = para.text.strip()
    
    # Основной шаблон: "Рисунок N - Текст" или "Рисунок N – Текст"
    pattern1 = r"^(Рисунок|Рис\.|Fig\.|Figure)\s+\d+\s*[-–]\s*.+$"
    
    # Альтернативный шаблон: "Рисунок N. Текст" или "Рис. N. Текст"
    pattern2 = r"^(Рисунок|Рис\.|Fig\.|Figure)\s+\d+\.\s*.+$"
    
    # Упрощенный шаблон: просто начинается с "Рисунок N" или "Рис. N"
    pattern3 = r"^(Рисунок|Рис\.|Fig\.|Figure)\s+\d+"
    
    # Проверяем все шаблоны
    if (re.match(pattern1, text, re.IGNORECASE) or 
        re.match(pattern2, text, re.IGNORECASE) or 
        re.match(pattern3, text, re.IGNORECASE)):
        
        # Дополнительные проверки для повышения точности
        
        # 1. Проверка длины текста (подписи обычно не очень длинные)
        if len(text) > 300:  # Слишком длинный для подписи
            return False
            
        # 2. Проверка выравнивания (обычно по центру)
        is_centered = False
        if hasattr(para, 'paragraph_format') and para.paragraph_format:
            if hasattr(para.paragraph_format, 'alignment') and para.paragraph_format.alignment:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                is_centered = para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
        
        # 3. Проверка на отсутствие других признаков (например, начало нового раздела)
        if check_all_runs_are_bold(para) and len(text) < 30:
            # Короткий жирный текст, начинающийся с "Рисунок" - скорее заголовок раздела
            return False
            
        # Если подпись выровнена по центру, это дополнительное подтверждение
        # Но мы все равно считаем это подписью к рисунку, даже если не выровнено
        return True
    
    return False

def is_table_title(para):
    """
    Проверяет, является ли параграф заголовком таблицы.
    
    Заголовок таблицы:
    1. Имеет формат "Таблица N - Текст" или "Таблица N – Текст" (с дефисом или тире)
    2. Может иметь стиль Caption или Table Caption или аналогичный
    3. Обычно размещается перед таблицей
    
    Args:
        para: Объект параграфа
    
    Returns:
        bool: True если параграф является заголовком таблицы
    """
    # Проверка по стилю
    if hasattr(para, 'style') and para.style:
        style_name = para.style.name.lower() if hasattr(para.style, 'name') else ""
        
        # Проверка на стили заголовков таблиц
        caption_indicators = ["caption", "подпись", "table", "таблица"]
        if any(indicator in style_name for indicator in caption_indicators):
            # Если это стиль подписи и текст содержит слово "таблица" - это заголовок таблицы
            if "таблица" in para.text.lower() or "table" in para.text.lower():
                return True
    
    # Проверка по формату текста
    pattern = r"^(Таблица|Табл\.|Tab\.|Table)\s+\d+\s*[-–]\s*.+$"
    text_matches = bool(re.match(pattern, para.text.strip(), re.IGNORECASE))
    
    # Если текст соответствует формату заголовка таблицы, это почти наверняка заголовок таблицы
    if text_matches:
        return True
    
    return False

def is_bibliography_item(para, in_bibliography_section):
    """Check if paragraph is a bibliography item."""
    # Если мы не в разделе библиографии, то с высокой вероятностью это не элемент библиографии
    if not in_bibliography_section:
        # Очень ограниченная проверка для случаев, когда раздел библиографии не был корректно определен
        text = para.text.strip()
        
        # Исключаем примечания и другие тексты
        if text.startswith("Примечание:") or text.startswith("Примечание "):
            return False
            
        # Проверка на встроенную нумерацию
        if hasattr(para, 'paragraph_format') and para.paragraph_format:
            if hasattr(para.paragraph_format, 'numbering') and para.paragraph_format.numbering:
                # Это встроенный список, но нужно дополнительно проверить, что это библиография
                if len(text) > 30 and ("//" in text or ": " in text or re.search(r"\d{4}\s*г", text)):
                    return True
            
        # Должно начинаться с цифры, точки и содержать специфические признаки библиографии
        # НО также не быть заголовком раздела или подраздела
        if (re.match(r"^\d+\.\s+", text) and 
            ("//" in text or ": " in text or re.search(r"\d{4}\s*г", text)) and
            not check_all_runs_are_bold(para) and
            len(text) > 50):
            return True
        return False
    
    # Если мы в разделе библиографии, более мягкая проверка
    text = para.text.strip()
    
    # Исключаем примечания и другие тексты, которые не являются библиографическими записями
    if text.startswith("Примечание:") or text.startswith("Примечание "):
        return False
    
    # Исключаем заголовки и подзаголовки
    if check_all_runs_are_bold(para):
        # Если текст жирный и короткий, вероятно это заголовок/подзаголовок
        if len(text) < 50:
            return False
    
    # Исключаем строки, которые похожи на объяснение оформления библиографии
    if "библиографические ссылки" in text.lower():
        return False
    if "оформление" in text.lower() and "источник" in text.lower():
        return False
    
    # Проверка на стиль библиографии
    if hasattr(para, 'style') and para.style:
        style_name = para.style.name.lower() if hasattr(para.style, 'name') else ""
        if 'bibliography' in style_name or 'источник' in style_name or 'reference' in style_name:
            return True
    
    # Проверка наличия встроенной нумерации (списков Word)
    if hasattr(para, 'paragraph_format') and para.paragraph_format:
        if hasattr(para.paragraph_format, 'numbering') and para.paragraph_format.numbering:
            # Это встроенный список, который может быть элементом библиографии
            # Но также нужно проверить его содержимое на библиографические признаки
            if len(text) > 30 and not check_all_runs_are_bold(para):
                # Дополнительная проверка на характерные признаки библиографии
                if ("//" in text or ": " in text or 
                    re.search(r"\d{4}", text) or 
                    "изд" in text.lower() or 
                    "с." in text):
                    return True
    
    # Проверка на формат "1. Автор..." - типичный для библиографии
    if re.match(r"^\d+\.\s+", text):
        # Это может быть библиографическая запись или заголовок
        is_bold = any(run.bold for run in para.runs if hasattr(run, 'bold') and run.bold)
        
        # Библиографические записи обычно не выделены жирным
        if not is_bold:
            return True
            
        # Проверка на библиографические особенности
        bibliography_indicators = ["//", ": ", ".: ", " под ред. ", " и др. ", "Т.", "Vol.", "№", "P.", "С."]
        for indicator in bibliography_indicators:
            if indicator in text:
                return True

    # Если текст длинный и содержит типичные элементы библиографии
    if len(text) > 50 and ("//" in text or ": " in text or re.search(r"\d{4}", text)):
        # Дополнительная проверка: параграф не должен быть заголовком
        if not check_all_runs_are_bold(para):
            return True
    
    return False

def is_appendix_heading(para):
    """Check if paragraph is an appendix heading."""
    # Проверяем независимо от регистра, удаляем точку в конце
    cleaned_text = para.text.strip().upper()
    if cleaned_text.endswith('.'):
        cleaned_text = cleaned_text[:-1]
    
    # Проверка на различные варианты заголовков приложений
    if cleaned_text.startswith("ПРИЛОЖЕНИЕ"):
        return True
        
    # Проверка по стилю
    if hasattr(para, 'style') and para.style:
        style_name = para.style.name.lower() if hasattr(para.style, 'name') else ""
        if "приложение" in style_name or "appendix" in style_name:
            return True
            
    return False

def is_paragraph_on_new_page(doc, paragraph_index):
    """Check if paragraph starts on a new page."""
    if paragraph_index <= 0:
        # First paragraph is always on a new page
        return True
    
    try:
        # Try to detect page breaks before the paragraph
        previous_para = doc.paragraphs[paragraph_index - 1]
        
        # Check for page break before via run properties
        if previous_para.runs:
            for run in previous_para.runs:
                if run.element.xpath("./w:br[@w:type='page']"):
                    return True
        
        # Check for paragraph properties
        if previous_para._p.xpath("./w:r/w:br[@w:type='page']"):
            return True
            
        # Check for section break with page break
        if previous_para._p.xpath("./w:pPr/w:sectPr"):
            return True
        
        return False
    except Exception:
        # In case of any error, assume it's not on a new page
        return False

def is_empty_paragraph(para):
    """Check if paragraph is empty (whitespace only)."""
    return not para.text.strip()

def has_spacing_after(para, next_para=None):
    """
    Check if paragraph has spacing after (either empty paragraph or spacing settings).
    
    Args:
        para: The paragraph to check
        next_para: The next paragraph if available
    
    Returns:
        bool: True if there is spacing after the paragraph
    """
    # Check for spacing after setting
    if hasattr(para, 'paragraph_format') and para.paragraph_format:
        if hasattr(para.paragraph_format, 'space_after') and para.paragraph_format.space_after:
            # Check if space_after is at least 6pt (0.5 line)
            if hasattr(para.paragraph_format.space_after, 'pt') and para.paragraph_format.space_after.pt >= 6:
                return True
    
    # Spacing between paragraphs is also applied if next paragraph has spacing before
    if next_para and hasattr(next_para, 'paragraph_format') and next_para.paragraph_format:
        if hasattr(next_para.paragraph_format, 'space_before') and next_para.paragraph_format.space_before:
            if hasattr(next_para.paragraph_format.space_before, 'pt') and next_para.paragraph_format.space_before.pt >= 6:
                return True
    
    return False

# --- Функции проверки форматирования ---

def check_font_formatting_for_runs(para, para_idx, comments_list, author, element_name,
                                   expected_font="Times New Roman", expected_size_pt=14,
                                   must_be_bold=False, expected_color_rgb=RGBColor(0,0,0)):
    """Общая функция для проверки шрифта, размера, жирности и цвета для всех runs абзаца."""
    if not para.runs and para.text.strip():
        comments_list.append((para_idx, f"Предупреждение ({element_name}): Не удалось проверить форматирование шрифта (отсутствуют 'runs' при наличии текста)", author))
        return

    # Собираем ошибки по всем runs, чтобы не дублировать сообщения для одного абзаца
    font_name_errors = set()
    font_size_errors = set()
    bold_errors = set()
    color_errors = set()

    for run in para.runs:
        if not run.text.strip(): continue

        # Шрифт
        font_name = get_run_font_name(run, para.style)
        if font_name and font_name != expected_font:
            font_name_errors.add(f"шрифт '{font_name}'")
        
        # Размер
        size_pt = get_run_font_size_pt(run, para.style)
        if size_pt is not None and abs(size_pt - expected_size_pt) > 0.1:
            font_size_errors.add(f"размер {size_pt:.0f}пт")

        # Жирность
        is_bold = get_run_bold_status(run, para.style)
        if is_bold is None: is_bold = False 
        if must_be_bold and not is_bold:
            bold_errors.add("не полужирный")

        # Цвет
        color_rgb = get_run_font_color_rgb(run, para.style)
        if color_rgb is not None and color_rgb != expected_color_rgb:
            color_errors.add(f"цвет {color_rgb}")
            
    # Формируем итоговое сообщение об ошибке, если есть
    final_errors = []
    if font_name_errors: final_errors.append(f"{', '.join(font_name_errors)} вместо '{expected_font}'")
    if font_size_errors: final_errors.append(f"{', '.join(font_size_errors)} вместо {expected_size_pt}пт")
    if bold_errors: final_errors.append(f"{', '.join(bold_errors)}")
    if color_errors: final_errors.append(f"{', '.join(color_errors)} вместо черного")

    if final_errors:
        comments_list.append((para_idx, f"Ошибка ({element_name}): {'; '.join(final_errors)}.", author))

def check_structural_or_appendix_heading_format(para, para_idx, comments_list, author, element_name):
    """Проверка для СТРУКТУРНЫХ заголовков и ПРИЛОЖЕНИЙ."""
    # Правила: 14 пт, черный, полужирный, по центру, без отступа первой строки
    check_font_formatting_for_runs(para, para_idx, comments_list, author, element_name, must_be_bold=True)

    alignment = get_effective_alignment(para)
    if alignment != WD_ALIGN_PARAGRAPH.CENTER:
        comments_list.append((para_idx, f"Ошибка ({element_name}): Выравнивание должно быть по центру (текущее: {alignment}).", author))

    first_line_indent_cm = get_first_line_indent_cm(para)
    if abs(first_line_indent_cm) > 0.01: # Отступ должен быть строго 0 (или очень близок к нему)
        comments_list.append((para_idx, f"Ошибка ({element_name}): Не должно быть отступа первой строки (текущий: {first_line_indent_cm:.2f} см).", author))
    
    # Точка в конце (только для простых заголовков без точки в самом названии)
    stripped_text_upper = para.text.strip().upper()
    # Проверяем, не является ли точка частью стандартного полного названия
    known_headings_with_possible_dot = ["СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ.", "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ."] # Маловероятно, но для примера
    if stripped_text_upper.endswith('.') and stripped_text_upper not in known_headings_with_possible_dot:
        # Дополнительная проверка, что это не "ПРИЛОЖЕНИЕ А."
        if not (element_name.startswith("Заголовок приложения") and stripped_text_upper.startswith("ПРИЛОЖЕНИЕ")):
             comments_list.append((para_idx, f"Ошибка ({element_name}): Не должно быть точки в конце.", author))

def check_page_margins(section, comments_list, author):
    """Check document margins."""
    # Expected margins in mm
    expected_left = 30
    expected_right = 15
    expected_top = 20
    expected_bottom = 20
    
    # Get actual margins in mm
    try:
        left_margin = section.left_margin.mm if section.left_margin else 0
        right_margin = section.right_margin.mm if section.right_margin else 0
        top_margin = section.top_margin.mm if section.top_margin else 0
        bottom_margin = section.bottom_margin.mm if section.bottom_margin else 0
        
        # Check each margin with some tolerance
        margin_errors = []
        if abs(left_margin - expected_left) > 1:
            margin_errors.append(f"левое поле (ожидается: {expected_left} мм, текущее: {left_margin:.1f} мм)")
        
        if abs(right_margin - expected_right) > 1:
            margin_errors.append(f"правое поле (ожидается: {expected_right} мм, текущее: {right_margin:.1f} мм)")
        
        if abs(top_margin - expected_top) > 1:
            margin_errors.append(f"верхнее поле (ожидается: {expected_top} мм, текущее: {top_margin:.1f} мм)")
        
        if abs(bottom_margin - expected_bottom) > 1:
            margin_errors.append(f"нижнее поле (ожидается: {expected_bottom} мм, текущее: {bottom_margin:.1f} мм)")
        
        # Add error message if any margins are incorrect
        if margin_errors:
            comments_list.append((-1, f"Ошибка: Неправильные поля страницы: {', '.join(margin_errors)}", author))
    except Exception as e:
        comments_list.append((-1, f"Ошибка при проверке полей страницы: {str(e)}", author))

def check_main_heading_format(para, para_idx, doc, comments_list, author, next_para=None):
    """Check main headings like ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ etc."""
    # Используем общую функцию для структурных заголовков
    check_structural_or_appendix_heading_format(para, para_idx, comments_list, author, "Основной заголовок")
    
    # Дополнительно проверяем отступ после заголовка
    if next_para and not is_empty_paragraph(next_para) and not has_spacing_after(para, next_para):
        comments_list.append((para_idx, f"Ошибка (Основной заголовок): После заголовка должен быть отступ (пустая строка или настройка интервала)", author))

def check_section_heading_format(para, para_idx, doc, comments_list, author, next_para=None):
    """Check formatting of section headings (1. Heading or 1 Heading)."""
    element_name = "Заголовок раздела"
    # Правила: 14 пт, черный, полужирный, по левому краю, отступ первой строки 1.25 см
    check_font_formatting_for_runs(para, para_idx, comments_list, author, element_name, must_be_bold=True)

    alignment = get_effective_alignment(para)
    if alignment != WD_ALIGN_PARAGRAPH.LEFT:
        comments_list.append((para_idx, f"Ошибка ({element_name}): Выравнивание должно быть по левому краю (текущее: {alignment}).", author))

    first_line_indent_cm = get_first_line_indent_cm(para)
    if abs(first_line_indent_cm - 1.25) > 0.1:
        comments_list.append((para_idx, f"Ошибка ({element_name}): Отступ первой строки должен быть 1.25 см (текущий: {first_line_indent_cm:.2f} см).", author))
    
    # Проверка формата номера "N." или "N "
    format_with_dot = re.match(r"^\d{1,2}\.\s+", para.text.strip())
    format_without_dot = re.match(r"^\d{1,2}\s+", para.text.strip())
    
    if not format_with_dot and not format_without_dot:
        comments_list.append((para_idx, f"Ошибка ({element_name}): Номер раздела должен быть в формате 'N. Название' или 'N Название', где N - число.", author))
    elif format_without_dot:
        # Предупреждение, если используется формат без точки после номера
        comments_list.append((para_idx, f"Предупреждение ({element_name}): Рекомендуется использовать формат 'N. Название' с точкой после номера.", author))
    
    # Точка в конце текстовой части заголовка
    text_content = para.text.strip()
    if format_with_dot:
        text_content = text_content[len(format_with_dot.group(0)):].strip()
    elif format_without_dot:
        text_content = text_content[len(format_without_dot.group(0)):].strip()
        
    if text_content.endswith('.'):
        comments_list.append((para_idx, f"Ошибка ({element_name}): Не должно быть точки в конце текстовой части заголовка.", author))

    # Проверка новой страницы
    if not is_paragraph_on_new_page(doc, para_idx):
        comments_list.append((para_idx, f"Ошибка ({element_name}): Заголовок раздела должен начинаться с новой страницы.", author))
        
    # Проверка отступа после заголовка
    if next_para and not is_empty_paragraph(next_para) and not has_spacing_after(para, next_para):
        comments_list.append((para_idx, f"Ошибка ({element_name}): После заголовка должен быть отступ (пустая строка или настройка интервала).", author))

def check_subsection_heading_format(para, para_idx, comments_list, author, next_para=None):
    """Check formatting of subsection headings (1.1 Heading without period)."""
    element_name = "Заголовок подраздела"
    # Правила: 14 пт, черный, полужирный, по левому краю, отступ первой строки 1.25 см
    check_font_formatting_for_runs(para, para_idx, comments_list, author, element_name, must_be_bold=True)

    alignment = get_effective_alignment(para)
    if alignment != WD_ALIGN_PARAGRAPH.LEFT:
        comments_list.append((para_idx, f"Ошибка ({element_name}): Выравнивание должно быть по левому краю (текущее: {alignment}).", author))

    first_line_indent_cm = get_first_line_indent_cm(para)
    if abs(first_line_indent_cm - 1.25) > 0.1:
        comments_list.append((para_idx, f"Ошибка ({element_name}): Отступ первой строки должен быть 1.25 см (текущий: {first_line_indent_cm:.2f} см).", author))
    
    # Проверка формата номера "N.M" (без точки в конце номера)
    format_correct = re.match(r"^(\d+(\.\d+)+)\s+", para.text.strip()) # Без точки в конце номера (правильно)
    format_incorrect = re.match(r"^(\d+(\.\d+)+)\.\s+", para.text.strip()) # С точкой в конце номера (неправильно)
    
    if not format_correct and not format_incorrect:
        comments_list.append((para_idx, f"Ошибка ({element_name}): Номер подраздела должен быть в формате 'N.M' (например, '1.1 Название'), без точки после номера.", author))
    elif format_incorrect:
        # Если найден формат с точкой после номера, это ошибка
        comments_list.append((para_idx, f"Ошибка ({element_name}): После номера подраздела (например, '{format_incorrect.group(1)}') не должно быть точки.", author))
    
    # Точка в конце текстовой части заголовка
    text_content = para.text.strip()
    if format_correct:
        text_content = text_content[len(format_correct.group(0)):].strip()
    elif format_incorrect:
        text_content = text_content[len(format_incorrect.group(0)):].strip()
        
    if text_content.endswith('.'):
        comments_list.append((para_idx, f"Ошибка ({element_name}): Не должно быть точки в конце текстовой части заголовка.", author))
    
    # Проверка отступа после заголовка
    if next_para and not is_empty_paragraph(next_para) and not has_spacing_after(para, next_para):
        comments_list.append((para_idx, f"Ошибка ({element_name}): После подзаголовка должен быть отступ (пустая строка или настройка интервала)", author))

def check_figure_caption_format(para, para_idx, comments_list, author):
    """Check formatting of figure captions."""
    # Check alignment (center)
    if hasattr(para, 'paragraph_format') and para.paragraph_format and hasattr(para.paragraph_format, 'alignment') and para.paragraph_format.alignment:
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            comments_list.append((para_idx, "Ошибка: Подпись к рисунку должна быть выровнена по центру", author))
    
    # Check font properties
    for run in para.runs:
        if hasattr(run, 'font') and run.font and hasattr(run.font, 'name') and run.font.name and run.font.name != "Times New Roman":
            comments_list.append((para_idx, f"Ошибка: Неправильный шрифт подписи к рисунку. Ожидается: Times New Roman. Текущий: {run.font.name}", author))
            break
        
        if hasattr(run, 'font') and run.font and hasattr(run.font, 'size') and run.font.size and run.font.size.pt != 14:
            comments_list.append((para_idx, f"Ошибка: Неправильный размер шрифта подписи к рисунку. Ожидается: 14 пт. Текущий: {run.font.size.pt} пт", author))
            break
            
        # Check font color
        if (hasattr(run, 'font') and run.font and hasattr(run.font, 'color') and 
            run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb):
            if run.font.color.rgb != RGBColor(0, 0, 0):
                comments_list.append((para_idx, f"Ошибка: Цвет шрифта подписи к рисунку должен быть черным", author))
                break
                    
    # Check format (Рисунок N – Title)
    pattern = r"^Рисунок\s+\d+\s*[-–]\s*.+$"
    if not re.match(pattern, para.text.strip()):
        comments_list.append((para_idx, "Ошибка: Неправильный формат подписи к рисунку. Должно быть 'Рисунок N – Название'", author))
    
    # Check period at end
    if para.text.strip().endswith('.'):
        comments_list.append((para_idx, "Ошибка: Подпись к рисунку не должна заканчиваться точкой", author))

def check_table_title_format(para, para_idx, comments_list, author):
    """Check formatting of table titles."""
    # Check alignment (left)
    if hasattr(para, 'paragraph_format') and para.paragraph_format and hasattr(para.paragraph_format, 'alignment') and para.paragraph_format.alignment:
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            comments_list.append((para_idx, "Ошибка: Заголовок таблицы должен быть выровнен по левому краю", author))
    
    # Check no first line indent
    if hasattr(para, 'paragraph_format') and para.paragraph_format and hasattr(para.paragraph_format, 'first_line_indent') and para.paragraph_format.first_line_indent:
        if para.paragraph_format.first_line_indent.cm > 0.1:
            comments_list.append((para_idx, f"Ошибка: У заголовка таблицы не должно быть отступа первой строки. Текущий: {para.paragraph_format.first_line_indent.cm:.2f} см", author))
    
    # Check font properties
    for run in para.runs:
        if hasattr(run, 'font') and run.font and hasattr(run.font, 'name') and run.font.name and run.font.name != "Times New Roman":
            comments_list.append((para_idx, f"Ошибка: Неправильный шрифт заголовка таблицы. Ожидается: Times New Roman. Текущий: {run.font.name}", author))
            break
                    
        if hasattr(run, 'font') and run.font and hasattr(run.font, 'size') and run.font.size and run.font.size.pt != 14:
            comments_list.append((para_idx, f"Ошибка: Неправильный размер шрифта заголовка таблицы. Ожидается: 14 пт. Текущий: {run.font.size.pt} пт", author))
            break
                    
        # Check font color
        if (hasattr(run, 'font') and run.font and hasattr(run.font, 'color') and 
            run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb):
            if run.font.color.rgb != RGBColor(0, 0, 0):
                comments_list.append((para_idx, f"Ошибка: Цвет шрифта заголовка таблицы должен быть черным", author))
            break
    
    # Check format (Таблица N – Title)
    pattern = r"^Таблица\s+\d+\s*[-–]\s*.+$"
    if not re.match(pattern, para.text.strip()):
        comments_list.append((para_idx, "Ошибка: Неправильный формат заголовка таблицы. Должно быть 'Таблица N – Название'", author))
    
    # Check period at end
    if para.text.strip().endswith('.'):
        comments_list.append((para_idx, "Ошибка: Заголовок таблицы не должен заканчиваться точкой", author))

def check_list_item_format(para, para_idx, comments_list, author, doc_paragraphs=None, current_para_idx=None):
    """Check formatting of list items."""
    # Проверка формата элемента списка
    text = para.text.strip()
    
    # Проверяем, предоставлены ли данные для проверки точки с запятой
    if doc_paragraphs is None or current_para_idx is None:
        doc_paragraphs = []
        current_para_idx = para_idx
    
    # Проверка на точку с запятой/точку в конце элемента списка
    next_para_idx = current_para_idx + 1
    is_next_para_list_item = False
    
    # Проверяем, является ли следующий параграф элементом списка
    if next_para_idx < len(doc_paragraphs):
        next_para = doc_paragraphs[next_para_idx]
        is_next_para_list_item = is_list_item(next_para)
    
    # Правило: если следующий параграф - элемент списка, то текущий должен заканчиваться точкой с запятой (;)
    # Если следующий параграф не элемент списка, то текущий должен заканчиваться точкой (.)
    if is_next_para_list_item:
        if not text.endswith(';'):
            comments_list.append((current_para_idx, "Ошибка: Элемент списка должен заканчиваться точкой с запятой (;), так как за ним следует другой элемент списка", author))
        if text.endswith('.'):
            comments_list.append((current_para_idx, "Ошибка: Элемент списка должен заканчиваться точкой с запятой (;), а не точкой, так как за ним следует другой элемент списка", author))
    else:
        # Последний элемент списка должен заканчиваться точкой
        if not text.endswith('.') and not text.endswith('!') and not text.endswith('?'):
            if text.endswith(';'):
                comments_list.append((current_para_idx, "Ошибка: Последний элемент списка должен заканчиваться точкой (.), а не точкой с запятой", author))
            else:
                comments_list.append((current_para_idx, "Ошибка: Элемент списка должен заканчиваться точкой (.)", author))
    
    # Проверяем, является ли элемент встроенным списком Word
    is_native_list = False
    is_numbered_list = False
    is_bulleted_list = False
    
    # Более подробная проверка встроенных списков
    style_name = "не определено"
    if hasattr(para, 'style') and para.style:
        if hasattr(para.style, 'name') and para.style.name:
            style_name = para.style.name.lower()
            if 'list' in style_name or 'numbering' in style_name or 'bullet' in style_name:
                is_native_list = True
                # Определяем тип списка по названию стиля
                if 'number' in style_name or 'numbering' in style_name:
                    is_numbered_list = True
                elif 'bullet' in style_name:
                    is_bulleted_list = True
    
    # Проверка по атрибутам нумерации
    numbering_info = "не найдена"
    has_numbering_attributes = False
    
    if hasattr(para, 'paragraph_format') and para.paragraph_format:
        if hasattr(para.paragraph_format, 'numbering') and para.paragraph_format.numbering:
            if para.paragraph_format.numbering.level is not None:
                is_native_list = True
                has_numbering_attributes = True
                numbering_info = f"level={para.paragraph_format.numbering.level}"
                
                # Определяем тип списка по формату нумерации
                if hasattr(para.paragraph_format.numbering, 'num_id') and para.paragraph_format.numbering.num_id is not None:
                    is_numbered_list = True
                    numbering_info += f", num_id={para.paragraph_format.numbering.num_id}"
                else:
                    is_bulleted_list = True
    
    # Очищаем текст от невидимых символов и пробелов в начале
    visible_text = ''.join(ch for ch in text if ch.isprintable()).lstrip()
    
    # Для встроенных списков Word проверяем наличие маркеров в тексте и стиль
    if is_native_list:
        # Если это параграф списка - проверяем маркер или его отсутствие
        try:
            # Проверяем наличие видимого маркера
            bullet_match = re.search(r'^[-•*]', visible_text)
            number_match = re.search(r'^\d+[.)]\s', visible_text)
            letter_match = re.search(r'^[а-яa-z][.)]\s', visible_text)
            
            # Определяем тип маркера для проверки
            has_bullet_marker = bullet_match is not None
            has_number_marker = number_match is not None
            has_letter_marker = letter_match is not None
            has_any_marker = has_bullet_marker or has_number_marker or has_letter_marker
            
            # Более точное определение типа списка на основе стиля и атрибутов нумерации
            # Сначала проверяем на основе стиля (если доступен)
            list_type = "unknown"
            list_type_source = "default"
            
            # Проверка по контексту - если это подпункт нумерованного списка
            # или предыдущий и следующий элементы нумерованные
            is_numbered_by_context = False
            context_info = []
            
            if current_para_idx > 0 and current_para_idx < len(doc_paragraphs) - 1:
                prev_para = doc_paragraphs[current_para_idx - 1]
                next_para = doc_paragraphs[current_para_idx + 1]
                
                # Проверяем видимые маркеры в соседних параграфах
                prev_has_number = False
                next_has_number = False
                
                if prev_para.text:
                    prev_has_number = bool(re.search(r'^\d+[.)]\s', prev_para.text.lstrip()))
                    prev_style = prev_para.style.name.lower() if prev_para.style and prev_para.style.name else "нет стиля"
                    context_info.append(f"Предыдущий параграф: стиль '{prev_style}', текст: '{prev_para.text[:20]}...'")
                    
                    # Проверка на нумерованный стиль в предыдущем параграфе
                    if "number" in prev_style or "enum" in prev_style:
                        prev_has_number = True
                        context_info.append("Предыдущий параграф имеет нумерованный стиль")
                
                if next_para.text:
                    next_has_number = bool(re.search(r'^\d+[.)]\s', next_para.text.lstrip()))
                    next_style = next_para.style.name.lower() if next_para.style and next_para.style.name else "нет стиля"
                    context_info.append(f"Следующий параграф: стиль '{next_style}', текст: '{next_para.text[:20]}...'")
                    
                    # Проверка на нумерованный стиль в следующем параграфе
                    if "number" in next_style or "enum" in next_style:
                        next_has_number = True
                        context_info.append("Следующий параграф имеет нумерованный стиль")
                
                # Дополнительная проверка на нумерованный стиль для текущего параграфа
                if style_name and ("list paragraph" in style_name.lower()):
                    # Если текущий параграф находится между нумерованными элементами,
                    # то это с высокой вероятностью тоже нумерованный элемент
                    if prev_has_number and next_has_number:
                        is_numbered_by_context = True
                        context_info.append("Элемент списка находится между нумерованными элементами")
                    elif prev_has_number:
                        # Если предыдущий элемент нумерованный, то этот скорее всего тоже
                        is_numbered_by_context = True
                        context_info.append("Элемент списка следует за нумерованным элементом")
                    elif next_has_number:
                        # Если следующий элемент нумерованный, то этот скорее всего тоже
                        is_numbered_by_context = True
                        context_info.append("Элемент списка предшествует нумерованному элементу")
                
                if prev_has_number or next_has_number:
                    is_numbered_by_context = True
                    context_info.append("Определен как нумерованный список по контексту")
            
            # Дополнительная проверка на наличие цифровых или буквенных маркеров в начале текста
            first_word = visible_text.split()[0] if visible_text.split() else ""
            if re.match(r"^\d+\.$", first_word) or re.match(r"^[а-яa-z]\.$", first_word, re.IGNORECASE):
                is_numbered_by_context = True
                context_info.append(f"Найден скрытый маркер нумерации: '{first_word}'")
            
            if style_name:
                if "bullet" in style_name.lower():
                    list_type = "bulleted"
                    list_type_source = "style_name_bullet"
                elif any(term in style_name.lower() for term in ["number", "enum", "numbered"]):
                    list_type = "numbered" 
                    list_type_source = "style_name_number"
            
            # Если тип не определен по стилю, проверяем атрибуты нумерации
            if list_type == "unknown" and hasattr(para, 'paragraph_format') and para.paragraph_format:
                if hasattr(para.paragraph_format, 'numbering') and para.paragraph_format.numbering:
                    if hasattr(para.paragraph_format.numbering, 'num_id') and para.paragraph_format.numbering.num_id:
                        list_type = "numbered"
                        list_type_source = "numbering_attr"
                    else:
                        list_type = "bulleted"
                        list_type_source = "default_numbering"
            
            # Если всё еще не определено, используем тип на основе маркеров в тексте
            if list_type == "unknown":
                if has_bullet_marker:
                    list_type = "bulleted"
                    list_type_source = "visible_marker_bullet"
                elif has_number_marker or has_letter_marker:
                    list_type = "numbered"
                    list_type_source = "visible_marker_number"
                elif is_numbered_by_context:
                    list_type = "numbered"
                    list_type_source = "context_analysis"
                elif is_bulleted_list:
                    list_type = "bulleted"
                    list_type_source = "is_bulleted_flag"
                elif is_numbered_list:
                    list_type = "numbered"
                    list_type_source = "is_numbered_flag"
                else:
                    # По умолчанию выбираем тип на основе контекста или предполагаем маркированный
                    if is_numbered_by_context:
                        list_type = "numbered"
                        list_type_source = "context_default"
                    else:
                        list_type = "bulleted"
                        list_type_source = "absolute_default"
            
            # Вывод отладочной информации
            print(f"--- Тип элемента списка: {list_type} (источник: {list_type_source})")
            if context_info:
                print("--- Контекстная информация:")
                for info in context_info:
                    print(f"    * {info}")
            
            # Для элементов между номерными списками явно устанавливаем тип на нумерованный
            if is_numbered_by_context and not has_any_marker and style_name == "list paragraph":
                list_type = "numbered"
                list_type_source = "context_override"
                print("    * ПРИМЕНЕНО: Переопределение типа на нумерованный на основе контекста")
            
            # Если это параграф списка без видимого маркера в тексте
            if style_name == "list paragraph" and not has_any_marker:
                # Определяем требуемый формат в зависимости от типа списка
                if list_type == "bulleted":
                    comments_list.append((current_para_idx, f"Ошибка: Элемент списка не содержит правильного маркера. "
                                         f"Для маркированного списка требуется маркер '- '", author))
                elif list_type == "numbered":
                    comments_list.append((current_para_idx, f"Ошибка: Элемент списка не содержит правильного маркера. "
                                         f"Для нумерованного списка требуется формат '1)' или 'а)' и т.п.", author))
                else:
                    comments_list.append((current_para_idx, f"Ошибка: Элемент списка не содержит правильного маркера. "
                                         f"Требуется: для маркированного списка - '- ', для нумерованного - '1)' или 'а)' и т.п.", author))
            else:
                # Для списков с видимыми маркерами проверяем соответствие типу
                if has_bullet_marker:
                    # Проверяем маркер маркированного списка
                    if not visible_text.startswith('- '):
                        comments_list.append((current_para_idx, "Ошибка: Неправильный маркер маркированного списка. "
                                             "Должен быть только маркер '- ' (дефис с пробелом)", author))
                elif has_number_marker or has_letter_marker:
                    # Проверяем маркер нумерованного списка
                    if not re.match(r"^\d+\)\s+", visible_text) and not re.match(r"^[а-яА-Я]\)\s+", visible_text):
                        comments_list.append((current_para_idx, "Ошибка: Неправильный маркер нумерованного списка. "
                                             "Допустимый формат: '1)' или 'а)' с пробелом после", author))
                # Если нет явного маркера, но это встроенный список - проверяем тип списка по атрибутам
                elif list_type == "bulleted":
                    comments_list.append((current_para_idx, "Ошибка: Элемент маркированного списка должен начинаться с '- ' (дефис с пробелом)", author))
                elif list_type == "numbered":
                    comments_list.append((current_para_idx, "Ошибка: Элемент нумерованного списка должен начинаться с '1)' или 'а)' с пробелом после", author))
        except Exception as e:
            # В случае ошибки при анализе маркеров добавим сообщение о возможной проблеме
            if style_name == "list paragraph":
                comments_list.append((current_para_idx, "Ошибка: Неправильный формат маркера элемента списка. "
                                     "Требуется: для маркированного списка - '- ', для нумерованного - '1)' или 'а)' и т.п.", author))
    else:
        # Для ручных списков проверка соответствия правильным форматам маркеров
        valid_markers = [
            r"^-\s+",         # Маркер дефис
            r"^\w\)\s+",      # Буквенный маркер с закрывающей скобкой
            r"^\d+\)\s+"      # Цифровой маркер с закрывающей скобкой
        ]
        
        # Проверка на неправильные форматы маркеров
        invalid_markers = [
            r"^\d+\.\s+",     # Цифра с точкой - неправильный формат
            r"^[а-яА-Яa-zA-Z]\.\s+.+$", # Буква с точкой (a., b., а. и т.д.) - неправильный формат
            r"^[*]\s+.+$",              # Звездочка (* ) - неправильный формат
            r"^[•●○◦■□▪▫]\s+.+$"        # Различные символы маркированного списка
        ]
        
        # Очищаем текст от невидимых символов и пробелов в начале
        visible_text = ''.join(ch for ch in text if ch.isprintable()).lstrip()
        
        # Основной критерий: начинается ли текст с допустимого маркера списка
        is_formatted_as_valid_list = any(bool(re.match(pattern, visible_text)) for pattern in valid_markers)
        
        # Проверяем, не является ли это неправильным форматом списка
        is_formatted_as_invalid_list = any(bool(re.match(pattern, visible_text)) for pattern in invalid_markers)
        
        # Определяем все элементы, которые соответствуют формату списка (правильные и неправильные)
        is_any_list_item = is_formatted_as_valid_list or is_formatted_as_invalid_list
        
        # Дополнительная проверка: не является ли это заголовком (не выделен жирным)
        if is_any_list_item:
            is_not_heading = not any(run.bold for run in para.runs if hasattr(run, 'bold') and run.bold)
            return is_not_heading
        
        return False

def check_bibliography_item_format(para, para_idx, comments_list, author):
    """
    Проверяет форматирование элемента библиографии (списка литературы).
    
    Элементы библиографии должны:
    - Иметь отступ первой строки 1.25 см
    - Выравнивание по ширине
    - Иметь правильный номер и формат
    
    Args:
        para: Объект параграфа
        para_idx: Индекс параграфа
        comments_list: Список для добавления комментариев
        author: Автор комментариев
    """
    # Пропускаем пустые параграфы
    if not para.text.strip():
        return
    
    # Дополнительная проверка, что это действительно элемент библиографии
    text = para.text.strip()
    
    # Пропускаем параграфы, которые описывают правила оформления библиографии
    if "библиографические ссылки" in text.lower():
        return
    
    # Пропускаем абзацы, содержащие ключевые слова заголовка библиографии
    text_upper = text.upper()
    for heading in ["СПИСОК", "ЛИТЕРАТУРА", "ИСТОЧНИК", "БИБЛИОГРАФИЯ"]:
        if heading in text_upper and len(text) < 50:
            return
    
    # Проверка наличия встроенного списка (нумерации)
    has_numbering = False
    numbering_format = None
    if hasattr(para, 'paragraph_format') and para.paragraph_format:
        if hasattr(para.paragraph_format, 'numbering') and para.paragraph_format.numbering:
            has_numbering = True
            # В этом случае нумерация обрабатывается Word автоматически
    
    # Проверка формата элемента библиографии
    # Если есть встроенная нумерация, считаем это правильным форматом
    # Если нет встроенной нумерации, должен начинаться с номера и точки
    if not has_numbering:
        bib_format_match = re.match(r"^(\d+)\.\s+(.+)$", text)
        if not bib_format_match:
            # Проверяем, что это не параграф с обычным текстом, который содержит библиографические ссылки
            if (not re.search(r"\[\d+\]", text) and 
                not "цитирования" in text.lower() and
                not "библиографическ" in text.lower()):
                comments_list.append((para_idx, "Ошибка: Элемент библиографии должен начинаться с номера и точки. Пример: '1. Иванов И.И.'", author))
        else:
            # Проверка содержимого записи на соответствие ГОСТ
            item_text = bib_format_match.group(2)
            check_gost_bibliography_compliance(item_text, para_idx, comments_list, author)
    else:
        # Если есть встроенная нумерация, проверяем только содержимое
        # Удаляем номер из начала текста, если он есть (иногда Word дублирует номер в тексте)
        clean_text = re.sub(r"^\d+[.\)]\s+", "", text)
        check_gost_bibliography_compliance(clean_text, para_idx, comments_list, author)
    
    # Проверка отступа первой строки с учетом стилей
    effective_indent = get_first_line_indent_cm(para)
    # Для встроенных списков отступ может быть другим из-за особенностей форматирования Word
    if has_numbering:
        # Для встроенных списков допускаем больший диапазон отступов
        if effective_indent < 0 or effective_indent > 2.5:
            comments_list.append((para_idx, f"Ошибка: Неправильный отступ первой строки элемента библиографии. Ожидается: 1.25 см. Текущий: {effective_indent:.2f} см", author))
    else:
        # Для обычных параграфов строгая проверка
        if abs(effective_indent - 1.25) > 0.1:
            comments_list.append((para_idx, f"Ошибка: Неправильный отступ первой строки элемента библиографии. Ожидается: 1.25 см. Текущий: {effective_indent:.2f} см", author))
    
    # Проверка выравнивания
    if hasattr(para, 'paragraph_format') and para.paragraph_format:
        effective_alignment = get_effective_alignment(para)
        if effective_alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            alignment_str = "не определено"
            if effective_alignment == WD_ALIGN_PARAGRAPH.LEFT:
                alignment_str = "по левому краю"
            elif effective_alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                alignment_str = "по правому краю"
            elif effective_alignment == WD_ALIGN_PARAGRAPH.CENTER:
                alignment_str = "по центру"
            
            comments_list.append((para_idx, f"Ошибка: Элемент библиографии должен быть выровнен по ширине, а не {alignment_str}", author))

def check_gost_bibliography_compliance(text, para_idx, comments_list, author):
    """
    Проверяет соответствие библиографической записи требованиям ГОСТ Р 7.0.100-2018.
    
    Основные правила ГОСТ Р 7.0.100-2018:
    1. Обязательные элементы: автор(ы), название, год издания
    2. Правильные разделители для разных типов источников
    3. Специфические особенности оформления разных типов источников
    
    Args:
        text: Текст библиографической записи
        para_idx: Индекс параграфа
        comments_list: Список для добавления комментариев
        author: Автор комментариев
    """
    # Удаляем номер из начала строки, если он есть
    if re.match(r"^\d+\.\s+", text):
        text = re.sub(r"^\d+\.\s+", "", text)
    
    # Пропускаем проверку, если текст не похож на библиографическую запись
    # (это может быть обычный текст в документе)
    if len(text) < 20:  # Слишком короткий для библиографической записи
        return
        
    # Пропускаем тексты, которые объясняют правила библиографии
    if "библиографические ссылки" in text.lower():
        return
    if "гост р 7.0." in text.lower() and len(text) < 50:
        return
        
    # Проверка на обязательные элементы ГОСТ
    
    # 1. Год издания (должен быть как минимум один год, обычно в конце)
    year_pattern = r"\b(19|20)\d{2}\b"
    if not re.search(year_pattern, text):
        comments_list.append((para_idx, "Ошибка: Отсутствует год издания в библиографической записи.", author))
    
    # 2. Название издания (должно быть выделено, обычно после автора)
    # Сложно проверить без доступа к форматированию отдельных частей текста
    
    # 3. Проверка разделителей
    # По ГОСТ используются // для разделения названия произведения и источника
    if ("//" not in text) and ("–" in text or "." in text) and len(text) > 50:
        # Это может быть книга или статья без указания журнала
        # Для книг // не требуется
        pass
    
    # 4. Проверка окончания
    # Библиографическая запись должна заканчиваться точкой
    if not text.strip().endswith("."):
        comments_list.append((para_idx, "Ошибка: Библиографическая запись должна заканчиваться точкой.", author))

def check_footnote_format(footnote, footnote_idx, comments_list, author):
    """Check formatting of footnotes."""
    # Check each paragraph in footnote
    try:
        for para_idx, para in enumerate(footnote.paragraphs):
            # Check font properties
            for run in para.runs:
                if hasattr(run, 'font') and run.font and hasattr(run.font, 'name') and run.font.name and run.font.name != "Times New Roman":
                    comments_list.append((-1, f"Ошибка: Неправильный шрифт сноски #{footnote_idx+1}. Ожидается: Times New Roman. Текущий: {run.font.name}", author))
                    break
                    
                # Check font size (10 or 12 pt)
                if hasattr(run, 'font') and run.font and hasattr(run.font, 'size') and run.font.size:
                    size_pt = run.font.size.pt
                    if size_pt < 10 or size_pt > 12:
                        comments_list.append((-1, f"Ошибка: Неправильный размер шрифта сноски #{footnote_idx+1}. Ожидается: 10-12 пт. Текущий: {size_pt} пт", author))
                    break
                    
                # Check font color
                if (hasattr(run, 'font') and run.font and hasattr(run.font, 'color') and 
                    run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb):
                    if run.font.color.rgb != RGBColor(0, 0, 0):
                        comments_list.append((-1, f"Ошибка: Цвет шрифта сноски #{footnote_idx+1} должен быть черным", author))
                    break
                    
            # Check line spacing (1.0, single)
            if hasattr(para, 'paragraph_format') and para.paragraph_format and hasattr(para.paragraph_format, 'line_spacing') and para.paragraph_format.line_spacing:
                if abs(para.paragraph_format.line_spacing - 1.0) > 0.01:
                    comments_list.append((-1, f"Ошибка: Неправильный межстрочный интервал сноски #{footnote_idx+1}. Ожидается: 1.0. Текущий: {para.paragraph_format.line_spacing}", author))
    except Exception as e:
        # Silently handle errors in footnote processing
        pass

def check_in_text_citations(doc_paragraphs, start_idx, comments_list, author):
    """Check in-text citations format."""
    citation_pattern = r'\[\d+(,\s*с\.\s*\d+)?\]'
    invalid_citation_pattern = r'\[\s+\d+|\d+\s+\]|\[\d+\s+,|\[\d+,\s+[^с]|\[\d+,\sс\s\.\s*\d+\]|\[\d+,\s*с\s+\.\s*\d+\]'
    
    for i, para in enumerate(doc_paragraphs):
        # Skip paragraphs before introduction
        if i < start_idx:
            continue
            
        # Skip if paragraph is empty
        if not para.text.strip():
            continue
        
        # Check for invalid citation formats
        invalid_citations = re.findall(invalid_citation_pattern, para.text)
        if invalid_citations:
            comments_list.append((i, f"Ошибка: Неправильный формат цитирования: {', '.join(invalid_citations)}. Должно быть [N] или [N, с. X]", author))

def check_appendix_heading_format(para, para_idx, doc, comments_list, author, next_para=None):
    """Check formatting of appendix headings (ПРИЛОЖЕНИЕ А)."""
    # Определяем ожидаемый формат
    expected_format = "ПРИЛОЖЕНИЕ " + para.text.strip().upper()[-1]
    if expected_format.endswith('.'):
        expected_format = expected_format[:-1]

    # Check new page
    if not is_paragraph_on_new_page(doc, para_idx):
        comments_list.append((para_idx, "Ошибка: Приложение должно начинаться с новой страницы", author))
    
    # Check alignment (center)
    if hasattr(para, 'paragraph_format') and para.paragraph_format and hasattr(para.paragraph_format, 'alignment'):
        if para.paragraph_format.alignment and para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            comments_list.append((para_idx, "Ошибка: Заголовок приложения должен быть выровнен по центру", author))
    
    # Check case (all uppercase)
    cleaned_text = para.text.strip()
    if cleaned_text.endswith('.'):
        cleaned_text = cleaned_text[:-1]
        
    if cleaned_text.upper() != cleaned_text:
        comments_list.append((para_idx, f"Ошибка: Заголовок приложения должен быть в верхнем регистре ({expected_format})", author))
    
    # Check first line indent (0 cm)
    if hasattr(para, 'paragraph_format') and para.paragraph_format and hasattr(para.paragraph_format, 'first_line_indent'):
        if para.paragraph_format.first_line_indent and para.paragraph_format.first_line_indent.cm > 0.1:
            comments_list.append((para_idx, "Ошибка: Заголовок приложения не должен иметь отступ первой строки", author))
    
    # Check font properties
    for run in para.runs:
        # Check bold
        if hasattr(run, 'bold') and not run.bold:
            comments_list.append((para_idx, "Ошибка: Заголовок приложения должен быть полужирным", author))
            break
        
        # Check font name
        if hasattr(run, 'font') and run.font and hasattr(run.font, 'name') and run.font.name and run.font.name != "Times New Roman":
            comments_list.append((para_idx, f"Ошибка: Неправильный шрифт заголовка приложения. Ожидается: Times New Roman. Текущий: {run.font.name}", author))
            break
        
        # Check font size
        if hasattr(run, 'font') and run.font and hasattr(run.font, 'size') and run.font.size and run.font.size.pt != 14:
            comments_list.append((para_idx, f"Ошибка: Неправильный размер шрифта заголовка приложения. Ожидается: 14 пт. Текущий: {run.font.size.pt} пт", author))
            break
    
    # Check period at end
    if para.text.strip().endswith('.'):
        comments_list.append((para_idx, "Ошибка: Заголовок приложения не должен заканчиваться точкой", author))
    
    # Check for spacing after heading
    if next_para and not is_empty_paragraph(next_para) and not has_spacing_after(para, next_para):
        comments_list.append((para_idx, "Ошибка: После заголовка приложения должен быть отступ (пустая строка или настройка интервала)", author))

def find_images_in_document(doc):
    """
    Находит все рисунки в документе и их позиции.
    
    Args:
        doc: документ docx
        
    Returns:
        list: список кортежей (paragraph_index, image_index)
    """
    images = []
    
    # Метод 1: Проверяем InlineShapes в каждом параграфе
    for i, para in enumerate(doc.paragraphs):
        has_image = False
        
        # Проверяем наличие InlineShapes в параграфе
        for run in para.runs:
            if hasattr(run, '_r') and run._r is not None:
                # Проверяем каждый run на наличие inline shape (XML-подход)
                if run._r.findall('.//wp:inline', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}):
                    has_image = True
                    break
                    
        # Также можем проверить напрямую элементы параграфа (другой подход)
        if hasattr(para, '_element') and para._element is not None:
            drawings = para._element.findall('.//w:drawing', 
                      namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if drawings:
                has_image = True
        
        if has_image:
            images.append((i, None))  # None вместо конкретного индекса изображения
    
    # Метод 2: Используем document.inline_shapes (если доступно)
    if hasattr(doc, 'inline_shapes') and doc.inline_shapes:
        # Для каждого изображения ищем, в каком параграфе оно находится
        for shape_idx, shape in enumerate(doc.inline_shapes):
            if shape.type == 3:  # WD_INLINE_SHAPE.PICTURE
                # К сожалению, нет прямого способа определить, в каком параграфе находится shape
                # Этот код только для реализации идеи, но реально он не выполнится
                pass
    
    # Добавляем отладочную информацию
    #print(f"DEBUG: Найдено {len(images)} изображений в документе")
    
    return images

def check_image_captions(doc, comments_list, author):
    """
    Проверяет соответствие рисунков и их подписей, последовательность нумерации.
    
    Args:
        doc: документ docx
        comments_list: список для добавления комментариев
        author: имя автора комментариев
    """
    # Найти все рисунки в документе
    images = find_images_in_document(doc)
    
    # Найти все подписи к рисункам
    captions = []
    
    # Используем функцию is_figure_caption для поиска подписей
    for i, para in enumerate(doc.paragraphs):
        if is_figure_caption(para):
            # Извлекаем номер рисунка из подписи
            text = para.text.strip()
            # Пробуем разные шаблоны для извлечения номера
            caption_num = None
            
            # Шаблон 1: "Рисунок N - Текст"
            match1 = re.match(r"^(?:Рисунок|Рис\.|Fig\.|Figure)\s+(\d+)\s*[-–]", text, re.IGNORECASE)
            if match1:
                caption_num = int(match1.group(1))
            
            # Шаблон 2: "Рисунок N. Текст"
            if caption_num is None:
                match2 = re.match(r"^(?:Рисунок|Рис\.|Fig\.|Figure)\s+(\d+)\.", text, re.IGNORECASE)
                if match2:
                    caption_num = int(match2.group(1))
            
            # Шаблон 3: "Рисунок N"
            if caption_num is None:
                match3 = re.match(r"^(?:Рисунок|Рис\.|Fig\.|Figure)\s+(\d+)", text, re.IGNORECASE)
                if match3:
                    caption_num = int(match3.group(1))
            
            if caption_num is not None:
                captions.append((i, caption_num, para))  # Сохраняем сам параграф для анализа выравнивания
                #print(f"DEBUG: Найдена подпись к рисунку {caption_num} в параграфе {i}: '{para.text}'")
    
    #print(f"DEBUG: Найдено {len(captions)} подписей к рисункам")
    
    # Добавить общий комментарий о количестве рисунков и подписей, независимо от их соответствия
    if len(images) > 0 or len(captions) > 0:
        comments_list.append((-1, f"Информация: В документе найдено {len(images)} рисунков и {len(captions)} подписей к рисункам.", author))
    
    # Проверить соответствие количества рисунков и подписей
    if len(images) != len(captions) and len(images) > 0:
        diff = abs(len(images) - len(captions))
        if len(images) > len(captions):
            comments_list.append((-1, f"Ошибка: В документе {len(images)} рисунков, но только {len(captions)} подписей. {diff} рисунок(ов) без подписи.", author))
        else:
            comments_list.append((-1, f"Ошибка: В документе {len(captions)} подписей к рисункам, но только {len(images)} рисунков. {diff} лишних подписей.", author))
    
    # Проверить последовательность нумерации
    if captions:
        expected_num = 1
        for i, num, _ in captions:
            if num != expected_num:
                comments_list.append((i, f"Ошибка: Нарушена последовательность нумерации рисунков. Ожидается: Рисунок {expected_num}, фактически: Рисунок {num}", author))
            expected_num += 1
    
    # Проверить, что после каждого рисунка следует подпись
    if images and captions:
        for img_idx, _ in images:
            caption_found = False
            nearest_caption_idx = -1
            nearest_caption_distance = float('inf')
            
            # Ищем ближайшую подпись после рисунка
            for caption_idx, caption_num, _ in captions:
                if caption_idx > img_idx and caption_idx - img_idx < nearest_caption_distance:
                    nearest_caption_distance = caption_idx - img_idx
                    nearest_caption_idx = caption_idx
                    if nearest_caption_distance <= 5:  # Увеличиваем максимальное расстояние до 5 параграфов
                        caption_found = True
                        break
                    
            if not caption_found:
                comments_list.append((img_idx, f"Ошибка: Рисунок в параграфе {img_idx} не имеет подписи или она расположена слишком далеко", author))
            elif nearest_caption_distance > 3:
                comments_list.append((img_idx, f"Предупреждение: Подпись к рисунку в параграфе {img_idx} расположена слишком далеко (через {nearest_caption_distance} параграфов)", author))
    
    # Проверить выравнивание параграфов с рисунками
    for img_idx, _ in images:
        para = doc.paragraphs[img_idx]
        if hasattr(para, 'paragraph_format') and hasattr(para.paragraph_format, 'alignment'):
            if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                actual_alignment = "по левому краю"
                if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    actual_alignment = "по правому краю"
                elif para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    actual_alignment = "по ширине"
                
                comments_list.append((img_idx, f"Ошибка: Рисунок должен быть выровнен по центру, а не {actual_alignment}", author))
        else:
            comments_list.append((img_idx, f"Ошибка: Рисунок должен быть выровнен по центру", author))
    
    # Проверить выравнивание подписей к рисункам
    for i, num, para in captions:
        if hasattr(para, 'paragraph_format') and hasattr(para.paragraph_format, 'alignment'):
            # Проверяем, что выравнивание по центру
            if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                actual_alignment = "по левому краю"
                if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    actual_alignment = "по правому краю"
                elif para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    actual_alignment = "по ширине"
                
                comments_list.append((i, f"Ошибка: Подпись к рисунку {num} должна быть выровнена по центру, а не {actual_alignment}", author))
        else:
            # Если не удалось определить выравнивание, предполагаем выравнивание по левому краю (по умолчанию)
            comments_list.append((i, f"Ошибка: Подпись к рисунку {num} должна быть выровнена по центру, а не по левому краю", author))

def find_tables_in_document(doc):
    """
    Находит все таблицы в документе и возвращает их индексы
    
    Args:
        doc: документ docx
        
    Returns:
        list: список кортежей (индекс элемента в документе, таблица)
    """
    tables = []
    
    # В docx таблицы находятся на верхнем уровне документа, поэтому нужно сложнее отслеживать их позицию
    # Проходим по всем элементам и отслеживаем текущий индекс параграфа
    current_idx = 0
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            tables.append((current_idx, element))
        elif element.tag.endswith('p'):
            current_idx += 1
    
    #print(f"DEBUG: Найдено {len(tables)} таблиц в документе")
    return tables

def check_table_captions(doc, comments_list, author):
    """
    Проверяет наличие и форматирование заголовков таблиц
    
    Args:
        doc: документ docx
        comments_list: список для добавления комментариев
        author: имя автора комментариев
    """
    # Найти все таблицы в документе
    tables = find_tables_in_document(doc)
    
    # Найти все заголовки таблиц
    captions = []
    caption_pattern = r"^Таблица\s+(\d+)\s*[-–]\s*.+$"
    for i, para in enumerate(doc.paragraphs):
        match = re.match(caption_pattern, para.text.strip())
        if match:
            caption_num = int(match.group(1))
            captions.append((i, caption_num, para))  # Сохраняем сам параграф для анализа выравнивания
            #print(f"DEBUG: Найден заголовок к таблице {caption_num} в параграфе {i}: '{para.text}'")
    
    #print(f"DEBUG: Найдено {len(captions)} заголовков таблиц")
    
    # Добавить общий комментарий о количестве таблиц и заголовков
    if len(tables) > 0 or len(captions) > 0:
        comments_list.append((-1, f"Информация: В документе найдено {len(tables)} таблиц и {len(captions)} заголовков к таблицам.", author))
    
    # Проверить соответствие количества таблиц и заголовков
    if len(tables) != len(captions) and len(tables) > 0:
        diff = abs(len(tables) - len(captions))
        if len(tables) > len(captions):
            comments_list.append((-1, f"Ошибка: В документе {len(tables)} таблиц, но только {len(captions)} заголовков. {diff} таблица(ц) без заголовка.", author))
        else:
            comments_list.append((-1, f"Ошибка: В документе {len(captions)} заголовков таблиц, но только {len(tables)} таблиц. {diff} лишних заголовков.", author))
    
    # Проверить последовательность нумерации таблиц
    expected_num = 1
    for i, num, _ in captions:
        if num != expected_num:
            comments_list.append((i, f"Ошибка: Нарушена последовательность нумерации таблиц. Ожидается: Таблица {expected_num}, фактически: Таблица {num}", author))
        expected_num += 1
    
    # Проверить, что перед каждой таблицей есть заголовок
    if tables and captions:
        for table_idx, _ in tables:
            caption_found = False
            nearest_caption_idx = -1
            
            # Ищем ближайший заголовок перед таблицей
            for caption_idx, caption_num, _ in captions:
                if caption_idx < table_idx and (nearest_caption_idx == -1 or caption_idx > nearest_caption_idx):
                    nearest_caption_idx = caption_idx
                    # Максимальное расстояние между заголовком и таблицей - 2 параграфа
                    if table_idx - caption_idx <= 2:
                        caption_found = True
            
            if not caption_found:
                if nearest_caption_idx != -1:
                    # Заголовок существует, но слишком далеко от таблицы
                    comments_list.append((table_idx, f"Ошибка: Заголовок таблицы должен быть размещен непосредственно перед таблицей (на расстоянии не более 1-2 параграфов)", author))
                else:
                    # Заголовок отсутствует
                    comments_list.append((table_idx, f"Ошибка: Таблица не имеет заголовка. Добавьте заголовок в формате 'Таблица N - Название таблицы'", author))
    
    # Проверить, что каждому заголовку соответствует таблица
    for caption_idx, caption_num, _ in captions:
        table_found = False
        
        # Ищем ближайшую таблицу после заголовка
        for table_idx, _ in tables:
            if table_idx > caption_idx and table_idx - caption_idx <= 2:
                table_found = True
                break
                    
        if not table_found:
            comments_list.append((caption_idx, f"Ошибка: Заголовок таблицы {caption_num} не соответствует ни одной таблице или таблица расположена слишком далеко", author))
    
    # Проверить выравнивание заголовков таблиц
    for i, num, para in captions:
        if hasattr(para, 'paragraph_format') and hasattr(para.paragraph_format, 'alignment'):
            # Проверяем только если выравнивание не по левому краю или None (так как None обычно означает по левому краю)
            if para.paragraph_format.alignment is not None and para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                actual_alignment = "по центру"
                if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    actual_alignment = "по правому краю"
                elif para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    actual_alignment = "по ширине"
                elif para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    actual_alignment = "по центру"
                
                comments_list.append((i, f"Ошибка: Заголовок таблицы {num} должен быть выровнен по левому краю, а не {actual_alignment}", author))
        else:
            # Если не удалось определить выравнивание, пропускаем сообщение об ошибке,
            # так как по умолчанию выравнивание обычно по левому краю
            pass

def check_bibliography_numbering(doc_paragraphs, bibliography_section_start, comments_list, author):
    """
    Проверяет правильность нумерации библиографических записей.
    
    Элементы библиографии должны:
    1. Иметь последовательную нумерацию (1, 2, 3, ...)
    2. Каждый элемент должен начинаться с номера и точки
    3. Не должно быть пропусков в нумерации
    
    Args:
        doc_paragraphs: Список параграфов документа
        bibliography_section_start: Индекс начала секции библиографии
        comments_list: Список для добавления комментариев
        author: Автор комментариев
    """
    if bibliography_section_start < 0:
        return
    
    # Находим все элементы библиографии
    bibliography_items = []
    in_bibliography_section = False
    found_real_bibliography_section = False
    
    for i, para in enumerate(doc_paragraphs):
        # Определяем начало и конец секции библиографии
        if is_bibliography_heading(para):
            in_bibliography_section = True
            found_real_bibliography_section = True
            continue
            
        # Если встретился другой структурный заголовок - выходим из секции библиографии
        if in_bibliography_section and is_main_heading(para) and not is_bibliography_heading(para):
            in_bibliography_section = False
            
        # Пропускаем случаи, когда абзац содержит ключевые слова заголовка библиографии,
        # но не является элементом библиографии
        text_upper = para.text.strip().upper()
        skip_paragraph = False
        for heading in ["СПИСОК", "ЛИТЕРАТУРА", "ИСТОЧНИК", "БИБЛИОГРАФИЯ"]:
            if heading in text_upper and len(text_upper) < 50:
                skip_paragraph = True
                break
        if skip_paragraph:
            continue
                
        # Добавляем только реальные элементы библиографии
        if in_bibliography_section and is_bibliography_item(para, True):
            # Проверка, является ли текущий абзац реальным элементом библиографии,
            # а не просто содержит библиографическую ссылку
            
            # Проверяем, есть ли встроенная нумерация
            has_numbering = False
            numbering_num = -1
            
            if hasattr(para, 'paragraph_format') and para.paragraph_format:
                if hasattr(para.paragraph_format, 'numbering') and para.paragraph_format.numbering:
                    has_numbering = True
                    # Для параграфов со встроенной нумерацией пытаемся определить номер
                    # Хотя python-docx не дает прямого доступа к номеру списка,
                    # мы можем попытаться его определить по количеству уже найденных элементов
                    # или по тексту, если он начинается с цифры (иногда Word дублирует номер в тексте)
                    text = para.text.strip()
                    number_match = re.match(r"^(\d+)[.\)]?\s+", text)
                    if number_match:
                        # Если в тексте есть номер, используем его
                        numbering_num = int(number_match.group(1))
                    else:
                        # Иначе предполагаем, что номер равен количеству уже найденных элементов + 1
                        numbering_num = len(bibliography_items) + 1
            
            # Если нет встроенной нумерации, пытаемся извлечь номер из текста
            if not has_numbering:
                number_match = re.match(r"^(\d+)\.", para.text.strip())
                if number_match:
                    numbering_num = int(number_match.group(1))
                    bibliography_items.append((i, numbering_num, para.text.strip(), has_numbering))
                else:
                    # Если нет явного номера, добавляем с номером -1 (что будет обработано как ошибка)
                    bibliography_items.append((i, -1, para.text.strip(), has_numbering))
            else:
                # Для параграфов со встроенной нумерацией добавляем с определенным номером
                bibliography_items.append((i, numbering_num, para.text.strip(), has_numbering))
    
    # Если не нашли ни одной записи или не нашли раздел библиографии, выходим
    if not bibliography_items or not found_real_bibliography_section:
        return
        
    # Проверяем правильность нумерации
    expected_number = 1
    
    for i, num, item_text, has_numbering in bibliography_items:
        if has_numbering:
            # Для встроенной нумерации проверяем, соответствует ли она ожидаемому номеру
            if num != expected_number:
                comments_list.append((i, f"Ошибка: Неправильная нумерация библиографической записи. "
                                  f"Ожидается: {expected_number}. Текущий: {num}", author))
            expected_number += 1
            continue
            
        if num == -1:
            # Элемент библиографии без номера - настоящая ошибка
            comments_list.append((i, "Ошибка: Библиографическая запись должна начинаться с номера и точки "
                               f"(ожидается: '{expected_number}. ')", author))
        elif num != expected_number:
            comments_list.append((i, f"Ошибка: Неправильная нумерация библиографической записи. "
                               f"Ожидается: {expected_number}. Текущий: {num}", author))
        
        expected_number += 1

def check_document_formatting_final(doc_path, author="Norm Control"):
    """
    Основная функция проверки форматирования документа
    
    Args:
        doc_path: путь к файлу docx
        author: имя автора, который будет указан в комментариях
        
    Returns:
        tuple: (список комментариев, путь к документу с комментариями)
    """
    try:
        doc = Document(doc_path)
        comments_to_add = []
        
        # Check page margins (applies to entire document)
        if doc.sections:
            check_page_margins(doc.sections[0], comments_to_add, author)
        
        # Process paragraphs
        processing_active = False
        in_bibliography_section = False
        intro_index = -1
        bibliography_index = -1  # Добавляем переменную для индекса начала библиографии
        
        for i, para in enumerate(doc.paragraphs):
            # Skip empty paragraphs
            if not para.text.strip():
                continue
            
            # Check if we've reached the ВВЕДЕНИЕ section
            if is_introduction_heading(para):
                processing_active = True
                intro_index = i
                next_para = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
                check_main_heading_format(para, i, doc, comments_to_add, author, next_para)
                continue
            
            # Check if we've reached the bibliography section
            if is_bibliography_heading(para):
                in_bibliography_section = True
                bibliography_index = i  # Устанавливаем индекс начала библиографии
                processing_active = True  # Ensure processing is active for bibliography
                next_para = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
                check_main_heading_format(para, i, doc, comments_to_add, author, next_para)
                continue
            
            # Skip formatting checks before ВВЕДЕНИЕ
            if not processing_active:
                continue
            
            # Check if paragraph is in a table
            if is_in_table(para, doc):
                continue  # Skip table content checks for now
            
            # Get the next paragraph for spacing checks if available
            next_para = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
            
            # Identify paragraph type and apply appropriate checks
            # Check appendix first (it has priority over main_heading)
            if is_appendix_heading(para):
                # Reset bibliography section flag if we've moved to appendices
                in_bibliography_section = False
                check_appendix_heading_format(para, i, doc, comments_to_add, author, next_para)
            elif is_main_heading(para):
                # Reset bibliography section flag if we've moved to another main section
                if in_bibliography_section and not is_bibliography_heading(para):
                    in_bibliography_section = False
                check_main_heading_format(para, i, doc, comments_to_add, author, next_para)
            elif is_section_heading(para):
                check_section_heading_format(para, i, doc, comments_to_add, author, next_para)
            elif is_subsection_heading(para):
                check_subsection_heading_format(para, i, comments_to_add, author, next_para)
            elif is_figure_caption(para):
                check_figure_caption_format(para, i, comments_to_add, author)
            elif is_table_title(para):
                check_table_title_format(para, i, comments_to_add, author)
            # Проверяем библиографические записи перед элементами списка, 
            # чтобы избежать ложных срабатываний для библиографических записей с номерами
            elif is_bibliography_item(para, in_bibliography_section):
                check_bibliography_item_format(para, i, comments_to_add, author)
            elif is_list_item(para) and not in_bibliography_section:  # Не проверяем элементы списка в библиографии
                check_list_item_format(para, i, comments_to_add, author, doc.paragraphs, i)
            else:
                # Assume it's regular main text
                check_main_text_format(para, i, comments_to_add, author)
        
        # Проверка соответствия рисунков и подписей
        check_image_captions(doc, comments_to_add, author)
        
        # Проверка соответствия таблиц и их заголовков
        check_table_captions(doc, comments_to_add, author)
        
        # Check footnotes if available
        try:
            if hasattr(doc.part.document, 'footnotes_part') and doc.part.document.footnotes_part:
                footnotes_part = doc.part.document.footnotes_part
                if hasattr(footnotes_part, 'footnotes') and footnotes_part.footnotes:
                    for idx, footnote_obj in enumerate(footnotes_part.footnotes.footnotes):
                        check_footnote_format(footnote_obj, idx, comments_to_add, author)
        except Exception as e:
            # Some documents might not have footnotes or the API might differ
            comments_to_add.append((-1, f"Предупреждение: Не удалось проверить сноски. {str(e)}", author))
        
        # Check in-text citations (only for paragraphs after ВВЕДЕНИЕ)
        if intro_index >= 0:
            check_in_text_citations(doc.paragraphs, intro_index, comments_to_add, author)
        
        # Проверка последовательности нумерации элементов библиографии
        check_bibliography_numbering(doc.paragraphs, bibliography_index, comments_to_add, author)
        
        return comments_to_add
    except Exception as e:
        # Return a meaningful error as a comment
        return [(0, f"Ошибка при проверке форматирования: {str(e)}", author)]

# Keep the original function for backwards compatibility
def check_document_formatting(doc_path, author="Norm Control"):
    """
    Legacy function for checking document formatting.
    
    Args:
        doc_path: path to the document
        author: name of the comment author (default "Norm Control")
        
    Returns:
        list: list of tuples (paragraph_index, comment_text, author)
        for detected formatting violations
    """
    return check_document_formatting_final(doc_path, author) 

def get_paragraph_type(para, doc, in_bibliography_section=False, previous_para_type=None):
    """
    Определение типа параграфа.
    
    Args:
        para: Объект параграфа
        doc: Объект документа
        in_bibliography_section: Флаг, находимся ли мы в разделе библиографии
        previous_para_type: Тип предыдущего параграфа, если известен
    
    Returns:
        str: Строка с типом параграфа
    """
    # Если параграф пустой, возвращаем "Пустой параграф"
    if not para.text.strip():
        return "Пустой параграф"
    
    # Проверка на элемент таблицы
    if is_in_table(para, doc):
        return "Элемент таблицы"
    
    # Проверки для подписей к рисункам и таблицам (они имеют высокий приоритет)
    # Подписи к рисункам могут быть в любом месте документа
    text = para.text.strip().lower()
    if "рисунок" in text or "рис." in text or "fig." in text or "figure" in text:
        if is_figure_caption(para):
            return "Подпись к рисунку"
    
    if "таблица" in text or "табл." in text or "table" in text or "tab." in text:
        if is_table_title(para):
            return "Заголовок таблицы"
        
    # Проверки, которые зависят от контекста (раздел библиографии)
    if in_bibliography_section:
        # В разделе библиографии элементы библиографии имеют приоритет над заголовками и списками
        if is_bibliography_item(para, in_bibliography_section):
            return "Элемент библиографии"
    
    # Проверки заголовков (от более специфичных к более общим)
    if is_main_heading(para):
        # Проверка на конкретные заголовки
        if is_bibliography_heading(para):
            return "Основной заголовок (БИБЛИОГРАФИЯ)"
        elif is_introduction_heading(para):
            return "Основной заголовок (ВВЕДЕНИЕ)"
        else:
            return "Основной заголовок"
    
    # Проверка на заголовки приложений (высокий приоритет)
    if is_appendix_heading(para):
        return "Заголовок приложения"
    
    # Повторная проверка для подписей к рисункам и таблицам с менее строгими критериями
    # Это поможет выявить подписи, которые не соответствуют точному формату
    if re.search(r"рисунок\s+\d+", text, re.IGNORECASE) or re.search(r"рис\.\s+\d+", text, re.IGNORECASE):
        # Дополнительные проверки, чтобы избежать ложных срабатываний
        # Например, проверка на длину текста и отсутствие признаков заголовка
        if len(text) < 300 and not check_all_runs_are_bold(para):
            return "Подпись к рисунку"
    
    if re.search(r"таблица\s+\d+", text, re.IGNORECASE) or re.search(r"табл\.\s+\d+", text, re.IGNORECASE):
        if len(text) < 300 and not check_all_runs_are_bold(para):
            return "Заголовок таблицы"
    
    # Проверка на элемент списка (перед проверкой заголовков разделов)
    # Это позволит избежать ложного определения элементов списка как заголовков
    if is_list_item(para):
        # Проверяем, не является ли это заголовком с форматом списка (1. Заголовок)
        # Заголовки обычно короткие и могут быть жирными
        text = para.text.strip()
        if re.match(r"^\d+\.\s+", text):
            is_bold = any(run.bold for run in para.runs if hasattr(run, 'bold') and run.bold)
            if is_bold and len(text) < 30:
                # Это может быть заголовок, проверим дальше
                pass
            else:
                # Это скорее список
                return "Элемент списка"
        else:
            # Это определенно список
            return "Элемент списка"
    
    # Проверки для заголовков разделов и подразделов
    # Если мы в разделе библиографии, проверяем сначала на элемент библиографии
    if not in_bibliography_section:
        if is_section_heading(para):
            return "Заголовок раздела"
            
        if is_subsection_heading(para):
            return "Подзаголовок"
    else:
        # В разделе библиографии текст с форматом "1. Текст" обычно является элементом библиографии
        # даже если выглядит как заголовок раздела
        text = para.text.strip()
        if re.match(r"^\d+\.\s+", text):
            return "Элемент библиографии"
    
    # Если не подошло ни к одному из специфических типов, считаем обычным текстом
    return "Обычный текст"

def is_list_item(para):
    """Check if paragraph is a list item."""
    # Текст параграфа
    text = para.text.strip()
    
    # Пустой параграф не может быть элементом списка
    if not text:
        return False
    
    # Проверка по стилю параграфа
    if hasattr(para, 'style') and para.style:
        if hasattr(para.style, 'name') and para.style.name:
            style_name = para.style.name.lower()
            if 'list' in style_name or 'numbering' in style_name or 'bullet' in style_name:
                return True
    
    # Проверка по атрибутам нумерации
    if hasattr(para, 'paragraph_format') and para.paragraph_format:
        if hasattr(para.paragraph_format, 'numbering') and para.paragraph_format.numbering:
            if para.paragraph_format.numbering.level is not None:
                return True
    
    # Очищаем текст от невидимых символов и пробелов в начале
    visible_text = text.lstrip()
    
    # Маркеры списка: дефис, буква+скобка, цифра+скобка
    list_markers = [
        r"^-\s+",         # Маркер дефис (-)
        r"^[а-яА-Яa-zA-Z]\)\s+",  # Буквенный маркер типа a), б) и т.д.
        r"^\d+\)\s+"      # Цифровой маркер типа 1), 2) и т.д.
    ]
    
    # Проверяем наличие маркера списка
    for marker in list_markers:
        if re.match(marker, visible_text):
            # Если это начинается с маркера списка - это элемент списка
            # Но нужно убедиться, что это не заголовок
            is_bold = any(run.bold for run in para.runs if hasattr(run, 'bold') and run.bold)
            
            # Если параграф не выделен жирным, скорее всего это список
            if not is_bold:
                return True
            else:
                # Если жирный, проверяем долю жирного текста
                bold_chars = 0
                total_chars = 0
                for run in para.runs:
                    if hasattr(run, 'text'):
                        run_text = run.text
                        total_chars += len(run_text)
                        if hasattr(run, 'bold') and run.bold:
                            bold_chars += len(run_text)
                
                # Если не весь текст жирный, то это может быть список с выделениями
                if total_chars > 0 and bold_chars / total_chars < 0.8:
                    return True
                    
    # Проверка на сомнительные случаи: цифра+точка
    number_dot_pattern = r"^\d+\.\s+.+"
    if re.match(number_dot_pattern, visible_text):
        # Это может быть элемент списка, заголовок или библиографическая запись
        
        # Если это заголовок (весь жирный) - не список
        is_bold = all(run.bold for run in para.runs if hasattr(run, 'bold'))
        if is_bold:
            return False
            
        # Если похоже на библиографическую запись - не список
        if "//" in visible_text or ".: " in visible_text:
            return False
            
        # Проверяем, похоже ли на заголовок по длине
        if len(visible_text) < 30:
            return False
            
        # Иначе это может быть элемент списка с неправильным форматом
        return True
    
    return False

def check_main_text_format(para, para_idx, comments_list, author):
    """
    Проверяет форматирование основного текста документа.
    
    Правила для основного текста:
    1. Шрифт: Times New Roman, 14 пт, черный
    2. Выравнивание: по ширине
    3. Отступ первой строки: 1.25 см
    4. Междустрочный интервал: 1.5
    
    Args:
        para: Объект параграфа
        para_idx: Индекс параграфа в документе
        comments_list: Список для добавления комментариев
        author: Автор комментариев
    """
    # Пропускаем пустые параграфы
    if not para.text.strip():
        return
    
    # Проверяем шрифт, размер и цвет для всех runs
    check_font_formatting_for_runs(
        para, para_idx, comments_list, author, "Основной текст", 
        expected_font="Times New Roman", expected_size_pt=14, 
        must_be_bold=False, expected_color_rgb=RGBColor(0,0,0)
    )
    
    # Проверяем выравнивание
    alignment = get_effective_alignment(para)
    if alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
        comments_list.append((para_idx, f"Ошибка (Основной текст): Выравнивание должно быть по ширине (текущее: {alignment}).", author))
    
    # Проверяем отступ первой строки
    first_line_indent_cm = get_first_line_indent_cm(para)
    if abs(first_line_indent_cm - 1.25) > 0.05:  # Допускаем небольшую погрешность
        comments_list.append((para_idx, f"Ошибка (Основной текст): Отступ первой строки должен быть 1.25 см (текущий: {first_line_indent_cm:.2f} см).", author))
    
    # Проверяем междустрочный интервал (если доступно)
    if hasattr(para, 'paragraph_format') and para.paragraph_format:
        if hasattr(para.paragraph_format, 'line_spacing') and para.paragraph_format.line_spacing:
            line_spacing = para.paragraph_format.line_spacing
            # Для междустрочного интервала 1.5 значение должно быть около 1.5
            if abs(line_spacing - 1.5) > 0.1:  # Допускаем небольшую погрешность
                comments_list.append((para_idx, f"Ошибка (Основной текст): Междустрочный интервал должен быть 1.5 (текущий: {line_spacing:.2f}).", author))