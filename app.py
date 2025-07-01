#!/usr/bin/env python3.13
# -*- coding: utf-8 -*-
"""
Приложение для проверки форматирования документов DOCX и добавления комментариев.
Требуется Python 3.13+ и python-docx 1.2.0+, которые поддерживают API для комментариев.
"""

from flask import Flask, render_template, request, send_from_directory, url_for, redirect, flash, jsonify
import os
import uuid
import time
import sys
import platform
from datetime import datetime
from werkzeug.utils import secure_filename
from pathlib import Path
import docx

# Импортируем существующие модули
from formatting_checker import check_document_formatting
from comment_utils import add_comments_to_docx

# Определяем базовую директорию приложения (для корректной работы абсолютных путей)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__)
app.secret_key = os.urandom(24)

# Настройки для загрузки файлов
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Ограничение размера файла 16MB

# Создаем директорию для загрузок, если она не существует
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    """Проверяет допустимое расширение файла"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_document_stats(file_path):
    """Получает статистику документа"""
    try:
        doc = docx.Document(file_path)
        stats = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': len(doc.sections),
            'pages': estimate_pages(doc)
        }
        return stats
    except Exception as e:
        print(f"Ошибка при получении статистики документа: {e}")
        return None

def estimate_pages(doc):
    """Здесь описана примерная оценка количества страниц в документе"""
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    # Добавляем символы из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_chars += sum(len(p.text) for p in cell.paragraphs)
    # Грубая эмпирическая оценка: ~3000 символов на страницу
    estimated_pages = max(1, round(total_chars / 3000))
    return estimated_pages

@app.route('/')
def index():
    """Главная страница с формой загрузки"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    """Обработка загрузки файла"""
    if 'docx_file' not in request.files:
        flash('Файл не выбран')
        return redirect(url_for('index'))
    
    file = request.files['docx_file']
    
    if file.filename == '':
        flash('Файл не выбран')
        return redirect(url_for('index'))
    
    if file and allowed_file(file.filename):
        # Безопасное сохранение файла
        filename = secure_filename(file.filename)
        unique_id = str(uuid.uuid4())
        temp_filename = f"{unique_id}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)
        file.save(file_path)
        
        
        author = request.form.get('author', 'Norm Control')
        output_prefix = request.form.get('output_prefix', '_with_remarks')
        
        try:
            
            document_stats = get_document_stats(file_path)
            
            
            comments = check_document_formatting(file_path, author)
            
            # Если есть комментарии, добавляем их в документ
            if comments:
                base_name = Path(filename).stem
                output_filename = f"{base_name}{output_prefix}.docx"
                output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                result_file = add_comments_to_docx(file_path, output_path, comments)
                
                
                return render_template('result.html', 
                                      filename=output_filename,
                                      comment_count=len(comments),
                                      document_stats=document_stats)
            else:
                
                return render_template('result.html', 
                                      no_comments=True,
                                      filename=filename,
                                      document_stats=document_stats)
                
        except Exception as e:
            flash(f"Ошибка при обработке файла: {e}")
            return redirect(url_for('index'))
            
    else:
        flash('Разрешены только файлы с расширением .docx')
        return redirect(url_for('index'))

@app.route('/download/<filename>')
def download(filename):
    """Скачивание обработанного файла"""
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/debug/paths')
def debug_paths():
    """Добавил этот маршрут для диагностики проблем на хостинге"""
    paths_info = {
        "BASE_DIR": BASE_DIR,
        "UPLOAD_FOLDER": UPLOAD_FOLDER,
        "upload_folder_exists": os.path.exists(UPLOAD_FOLDER),
        "upload_folder_is_dir": os.path.isdir(UPLOAD_FOLDER),
        "upload_folder_writable": os.access(UPLOAD_FOLDER, os.W_OK),
        "uploaded_files": os.listdir(UPLOAD_FOLDER) if os.path.exists(UPLOAD_FOLDER) else []
    }
    return jsonify(paths_info)

@app.route('/debug/test-upload', methods=['GET', 'POST'])
def test_upload():
    """Тестовый маршрут для проверки загрузки файлов"""
    if request.method == 'GET':
        return '''
        <!doctype html>
        <title>Тест загрузки файлов</title>
        <h1>Загрузите тестовый файл</h1>
        <form method=post enctype=multipart/form-data>
          <input type=file name=file>
          <input type=submit value=Загрузить>
        </form>
        '''
    
    if 'file' not in request.files:
        return 'Файл не выбран'
    
    file = request.files['file']
    if file.filename == '':
        return 'Файл не выбран'
    
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    return f'''
    <h1>Файл успешно загружен</h1>
    <p>Путь: {file_path}</p>
    <p>Размер: {os.path.getsize(file_path)} байт</p>
    <p>Скачать: <a href="/download/{filename}">{filename}</a></p>
    '''

@app.route('/debug')
def debug():
    """Страница диагностики приложения"""
    # Информация о системе
    try:
        import flask
        flask_version = flask.__version__
    except ImportError:
        flask_version = "Не установлен"
    
    system_info = {
        "working_directory": os.getcwd(),
        "upload_folder": UPLOAD_FOLDER,
        "python_version": platform.python_version(),
        "flask_version": flask_version
    }
    
    # Проверка директории загрузок
    upload_exists = os.path.exists(UPLOAD_FOLDER)
    upload_writable = os.access(UPLOAD_FOLDER, os.W_OK) if upload_exists else False
    
    files = []
    file_count = 0
    
    if upload_exists:
        try:
            file_list = os.listdir(UPLOAD_FOLDER)
            file_count = len(file_list)
            
            for file_name in file_list[:10]:  # Показываем только первые 10 файлов
                file_path = os.path.join(UPLOAD_FOLDER, file_name)
                if os.path.isfile(file_path):
                    size_kb = round(os.path.getsize(file_path) / 1024, 2)
                    mod_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
                    files.append({"name": file_name, "size": size_kb, "modified": mod_time})
        except Exception as e:
            files = [{"name": f"Ошибка чтения директории: {str(e)}", "size": 0, "modified": ""}]
    
    upload_check = {
        "exists": upload_exists,
        "writable": upload_writable,
        "file_count": file_count,
        "files": files
    }
    
    # Проверка наличия модулей
    modules = [
        {"name": "formatting_checker.py", "exists": os.path.exists(os.path.join(BASE_DIR, "formatting_checker.py"))},
        {"name": "comment_utils.py", "exists": os.path.exists(os.path.join(BASE_DIR, "comment_utils.py"))}
    ]
    
    return render_template('debug.html', 
                          system_info=system_info, 
                          upload_check=upload_check,
                          modules_check=modules)

def cleanup_old_files(max_age_hours=24):
    """Сделал функцию для очистки старых файлов, чтобы не захламлять сервер"""
    current_time = time.time()
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.isfile(file_path) and os.path.exists(file_path):
            file_age_hours = (current_time - os.path.getmtime(file_path)) / 3600
            if file_age_hours > max_age_hours:
                try:
                    os.remove(file_path)
                except Exception:
                    pass

if __name__ == '__main__':
    # При запуске приложения очищаем старые файлы
    cleanup_old_files()
    # Для локального запуска используем debug=True
    app.run(debug=True)
    
# На PythonAnywhere будет использоваться application = app 
application = app 