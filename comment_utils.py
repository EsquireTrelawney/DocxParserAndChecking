import uuid
from datetime import datetime
from lxml import etree
from docx import Document
import zipfile
import os
import shutil
from pathlib import Path

def qn(tag):
    """
    Я сделал эту функцию для квалификации имен XML тегов в docx
    """
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    prefix, tagroot = tag.split(':')
    return '{{{}}}{}'.format(namespace[prefix], tagroot)

def is_paragraph_in_table(paragraph_element, namespaces):
    """
    Проверяет, находится ли параграф внутри таблицы.
    
    Args:
        paragraph_element: XML-элемент параграфа
        namespaces: словарь namespaces для xpath
    
    Returns:
        bool: True, если параграф находится внутри таблицы
    """
    parent = paragraph_element.getparent()
    while parent is not None:
        if parent.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc':
            return True
        parent = parent.getparent()
    return False

def add_comments_to_docx(input_path, output_path, comments_info):
    """
    Добавляет комментарии в DOCX документ
    
    Args:
        input_path: путь к исходному документу
        output_path: путь для сохранения документа с комментариями
        comments_info: список кортежей (paragraph_index, comment_text, author)
    """
    
    doc = Document(input_path)
    
   
    debug_info = []
    debug_info.append(f"Всего параграфов в документе: {len(doc.paragraphs)}")
    
    # Сортируем комментарии по возрастанию индекса параграфа
    sorted_comments = sorted(comments_info, key=lambda x: x[0] if x[0] >= 0 else float('inf'))
    
    # А это - словарь для хранения отображений специальных индексов на реальные параграфы
    special_index_mapping = {
        -1: 0,  # Индекс -1 (сноски) -> первый параграф
    }
    
    # Фильтруем параграфы, исключая те, которые находятся в таблицах
    body_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        # Проверяем, не находится ли параграф в таблице
        in_table = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if para in cell.paragraphs:
                        in_table = True
                        break
                if in_table:
                    break
            if in_table:
                break
        
        if not in_table:
            body_paragraphs.append(para)
    
    debug_info.append(f"Параграфов основного тела (не в таблицах): {len(body_paragraphs)}")
    
    # Добавляем комментарии
    for comment_index, (paragraph_index, comment_text, author) in enumerate(sorted_comments):
        if paragraph_index < 0:
            if paragraph_index in special_index_mapping:
                paragraph_index = special_index_mapping[paragraph_index]
            else:
                # Для неизвестных отрицательных индексов используем первый параграф
                paragraph_index = 0
            
            # Добавляем префикс к тексту комментария для общих комментариев
            if "[Общий комментарий] " not in comment_text:
                comment_text = f"[Общий комментарий] {comment_text}"
                
            
            target_para = doc.paragraphs[paragraph_index]
        else:
            
            # Проверка на валидность индекса
            if paragraph_index >= len(body_paragraphs):
                debug_info.append(f"Ошибка: Индекс {paragraph_index} вне диапазона основного тела (всего {len(body_paragraphs)} параграфов)")
                # Используем последний параграф основного тела, если индекс вне диапазона
                paragraph_index = len(body_paragraphs) - 1
            
            
            target_para = body_paragraphs[paragraph_index]
        
        # Добавляю ID параграфа для отладки
        debug_id = f"P{paragraph_index}_C{comment_index}"
        
        # Добавляею информацию об отладке в текст комментария
        full_comment_text = f"{comment_text} [Debug ID: {debug_id}]"
        
        # Создаю здесь комментарий с использованием метода add_comment
        try:
            # Получаем все runs в параграфе
            runs = target_para.runs
            
            if runs:
                
                comment = doc.add_comment(runs, full_comment_text, author)
                debug_info.append(f"Комментарий добавлен: {debug_id} к параграфу {paragraph_index} через существующие runs")
            else:
                # Пришлось сделать так: если runs нет, создаем новый run
                run = target_para.add_run(" ")
                comment = doc.add_comment(run, full_comment_text, author)
                debug_info.append(f"Комментарий добавлен: {debug_id} к параграфу {paragraph_index} через новый run")
        except Exception as e:
            # Если что-то пошло не так, добавляем информацию об ошибке в отладочный отчет
            debug_info.append(f"Ошибка при добавлении комментария {debug_id}: {str(e)}")
    
    
    doc.save(output_path)
    
    # Для создания отладочного файл
    debug_path = output_path + '.debug.txt'
    with open(debug_path, 'w', encoding='utf-8') as debug_file:
        debug_file.write('\n'.join(debug_info))
    
    return output_path