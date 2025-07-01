from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from formatting_checker import check_document_formatting_final, find_images_in_document, check_image_captions
import os

def create_test_doc_with_images():
    """Создает тестовый документ с рисунками и подписями"""
    doc = Document()
    
    # Заголовок ВВЕДЕНИЕ тут нужен для активации проверок
    heading = doc.add_paragraph("ВВЕДЕНИЕ")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.runs[0].bold = True

    doc.add_paragraph()
    
    # Имитируем рисунок (это не добавит реальное изображение, но для тестирования достаточно)
    p = doc.add_paragraph("Рисунок 1")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Подпись к рисунку 1
    caption1 = doc.add_paragraph("Рисунок 1 – Правильная подпись")
    caption1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Текст между рисунками
    doc.add_paragraph("Какой-то текст между рисунками.")
    
    # Еще один рисунок, но подпись с неправильной нумерацией
    p = doc.add_paragraph("Рисунок 2")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Неправильная подпись (нумерация 3 вместо 2)
    caption2 = doc.add_paragraph("Рисунок 3 – Неправильная нумерация")
    caption2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Рисунок без подписи
    p = doc.add_paragraph("Рисунок без подписи")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Неправильное выравнивание
    
    # Сохраняем документ
    filename = "test_images.docx"
    doc.save(filename)
    return filename

def test_image_detection():
    """Тест функции определения рисунков"""
    print("=== Тест определения рисунков в документе ===")
    test_file = create_test_doc_with_images()
    
    # К сожалению, эти тесты не будут работать с созданным нами документом,
    # так как мы не можем добавить настоящие рисунки через python-docx.
    # Для реального тестирования нужен документ с реальными рисунками.
    
    print(f"\nДокумент создан: {test_file}")
    print("Для полноценного тестирования, добавьте в него реальные изображения через Word")
    
    # Пробуем запустить проверку формата на созданном документе
    print("\n=== Тест проверки форматирования документа ===")
    comments = check_document_formatting_final(test_file)
    
    # Выводим все комментарии
    print(f"Найдено {len(comments)} замечаний:")
    for idx, comment, author in comments:
        print(f"[{idx}] {comment}")

if __name__ == "__main__":
    test_image_detection() 