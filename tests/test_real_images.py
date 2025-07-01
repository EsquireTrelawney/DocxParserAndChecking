import sys
import os
from docx import Document
from formatting_checker import find_images_in_document, check_image_captions, check_document_formatting_final

def detect_images(docx_path):
    """
    Анализирует документ и пытается найти в нем изображения разными способами.
    """
    print(f"\nАнализ документа на наличие изображений: {docx_path}")
    
    if not os.path.exists(docx_path):
        print(f"Ошибка: Файл {docx_path} не найден!")
        return
    
    try:
        doc = Document(docx_path)
        
        # Здесь использован прямой доступ к document.inline_shapes
        if hasattr(doc, 'inline_shapes'):
            print(f"\nМетод 1 - doc.inline_shapes:")
            print(f"Найдено {len(doc.inline_shapes)} встроенных объектов")
            
            for i, shape in enumerate(doc.inline_shapes):
                shape_type = "Неизвестно"
                if hasattr(shape, 'type'):
                    if shape.type == 3:  # WD_INLINE_SHAPE.PICTURE
                        shape_type = "Изображение"
                    elif shape.type == 4:  # WD_INLINE_SHAPE.CHART
                        shape_type = "Диаграмма"
                    elif shape.type == 8:  # WD_INLINE_SHAPE.LINKED_PICTURE
                        shape_type = "Связанное изображение"
                print(f"  Объект {i+1}: Тип={shape_type}")
        
        # Или опционально попытка использовать мою функци.
        print("\nМетод 2 - Наша функция find_images_in_document:")
        images = find_images_in_document(doc)
        print(f"Найдено {len(images)} изображений")
        for i, (para_idx, _) in enumerate(images):
            print(f"  Изображение {i+1}: В параграфе {para_idx}, начало текста: '{doc.paragraphs[para_idx].text[:30]}...'")
        
        # И наконец попытка XML-анализ всего документа
        print("\nМетод 3 - Прямой XML-анализ:")
        
        # Проверяем все параграфы на наличие элементов рисунков через XML
        drawing_count = 0
        picture_count = 0
        
        for i, para in enumerate(doc.paragraphs):
            if not hasattr(para, '_element'):
                continue
                
            drawings = para._element.findall('.//w:drawing', 
                      namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            pics = para._element.findall('.//pic:pic', 
                   namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'})
            
            if drawings:
                drawing_count += len(drawings)
                print(f"  Параграф {i}: {len(drawings)} w:drawing элементов")
            
            if pics:
                picture_count += len(pics)
                print(f"  Параграф {i}: {len(pics)} pic:pic элементов")
                
        print(f"Всего найдено чистым XML-анализом: {drawing_count} drawing элементов, {picture_count} picture элементов")
        
        # Анализ подписей к рисункам
        print("\nАнализ подписей к рисункам:")
        comments_list = []
        check_image_captions(doc, comments_list, "Тест")
        
        print("\nРезультаты проверки:")
        for idx, comment, _ in comments_list:
            print(f"[{idx}] {comment}")
            
    except Exception as e:
        print(f"Ошибка при анализе документа: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    # Используем конкретный файл вместо запроса пользователя
    docx_path = "test_normcontrol_documentFULL.docx"
    
    if not os.path.exists(docx_path):
        print(f"Файл не найден: {docx_path}")
        # Резервный вариант, если файл не найден
        docx_path = "test_document_for_normcontrol.docx" 
        if not os.path.exists(docx_path):
            print(f"Резервный файл также не найден: {docx_path}")
            sys.exit(1)
    
    detect_images(docx_path) 