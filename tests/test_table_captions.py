#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
from docx import Document
from formatting_checker import find_tables_in_document, check_table_captions

def test_table_detection():
    """Тестирование обнаружения таблиц в документе"""
    # Используем тестовый документ с таблицами
    doc_path = 'test_normcontrol_documentFULL.docx'
    
    # Загружаем документ
    doc = Document(doc_path)
    
    # Проверяем обнаружение таблиц
    tables = find_tables_in_document(doc)
    
    print(f"\nАнализ документа на наличие таблиц: {doc_path}\n")
    print(f"Метод 1 - Наша функция find_tables_in_document:")
    print(f"Найдено {len(tables)} таблиц")
    
    # Выводим информацию о найденных таблицах
    for i, (table_idx, _) in enumerate(tables):
        print(f"  Таблица {i+1}: Расположена перед параграфом {table_idx}")
    
    # Метод 2 - Стандартная функция python-docx
    print(f"\nМетод 2 - doc.tables:")
    print(f"Найдено {len(doc.tables)} таблиц")
    
    return tables

def test_table_captions():
    """Тестирование проверки заголовков таблиц"""
    # Используем тестовый документ с таблицами
    doc_path = 'test_normcontrol_documentFULL.docx'
    
    # Загружаем документ
    doc = Document(doc_path)
    
    # Список для хранения комментариев
    comments_list = []
    
    # Проверяем заголовки таблиц
    print(f"\nАнализ заголовков таблиц в документе: {doc_path}\n")
    check_table_captions(doc, comments_list, "Test")
    
    # Выводим обнаруженные проблемы
    for para_idx, comment, author in comments_list:
        print(f"[{para_idx}] {comment}")
    
    return comments_list

if __name__ == "__main__":
    print("\nТестирование обнаружения таблиц в документе")
    tables = test_table_detection()
    
    print("\nТестирование проверки заголовков таблиц")
    comments = test_table_captions() 