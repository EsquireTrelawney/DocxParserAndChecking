#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
import time
import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from formatting_checker import is_list_item, check_list_item_format

def create_test_doc_with_native_lists():
    """Создает тестовый документ с встроенными списками Word"""
    doc = Document()
    
    # Добавляем заголовок ВВЕДЕНИЕ для активации проверок
    heading = doc.add_paragraph("ВВЕДЕНИЕ")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.runs[0].bold = True
    
    # Пустая строка
    doc.add_paragraph()
    
    # Добавляем обычный текст
    doc.add_paragraph("Это обычный текст параграфа, который не является элементом списка.")
    
    # Заголовок для раздела со списками
    subsection = doc.add_paragraph("Тестовые встроенные списки")
    subsection.runs[0].bold = True
    
    # ====== Правильно отформатированные списки ======
    doc.add_paragraph("Маркированный список с правильным форматом (дефис):")
    
    # Метод создания списка с помощью стиля "List Bullet"
    for item in ["Первый элемент правильного маркированного списка (стиль);", 
                "Второй элемент правильного маркированного списка (стиль);",
                "Третий элемент правильного маркированного списка (стиль)."]:
        p = doc.add_paragraph("- " + item, style='List Bullet')
        
    # Встроенный нумерованный список
    doc.add_paragraph()
    doc.add_paragraph("Нумерованный список с правильным форматом (цифра со скобкой):")
    
    # Создание нумерованного списка со скобками
    for i, item in enumerate(["Первый элемент правильного нумерованного списка (стиль);", 
                        "Второй элемент правильного нумерованного списка (стиль);",
                        "Третий элемент правильного нумерованного списка (стиль)."]):
        p = doc.add_paragraph(f"{i+1}) " + item, style='List Number')
    
    # ====== Неправильно отформатированные списки ======
    doc.add_paragraph()
    doc.add_paragraph("Маркированный список с неправильным форматом (стандартная точка Word):")
    
    # Стандартный встроенный маркированный список Word с точками
    for item in ["Первый элемент неправильного маркированного списка;", 
                "Второй элемент неправильного маркированного списка;",
                "Третий элемент неправильного маркированного списка."]:
        p = doc.add_paragraph(item, style='List Bullet')
        # Удаляем "- " с начала, чтобы симулировать маркер "•"
        if p.runs and p.runs[0].text.startswith("- "):
            p.runs[0].text = p.runs[0].text[2:]
    
    # Встроенный нумерованный список с неправильным форматом
    doc.add_paragraph()
    doc.add_paragraph("Нумерованный список с неправильным форматом (стандартный формат с точкой):")
    
    # Стандартный нумерованный список с нумерацией "1."
    for i, item in enumerate(["Первый элемент неправильного нумерованного списка;", 
                        "Второй элемент неправильного нумерованного списка;",
                        "Третий элемент неправильного нумерованного списка."]):
        p = doc.add_paragraph(item, style='List Number')
        # Удаляем скобку и заменяем на точку
        if p.runs and p.runs[0].text.startswith(f"{i+1})"):
            p.runs[0].text = p.runs[0].text.replace(f"{i+1})", f"{i+1}.")
            
    # ====== Тест определения по контексту ======
    doc.add_paragraph()
    doc.add_paragraph("Нумерованный список с элементами без маркеров (определение по контексту):")
    
    # Создаем нумерованный список, где часть элементов без маркеров
    p = doc.add_paragraph("1) Первый элемент нумерованного списка с маркером;", style='List Number')
    p = doc.add_paragraph("Второй элемент нумерованного списка без маркера;", style='List Paragraph')
    p = doc.add_paragraph("3) Третий элемент нумерованного списка с маркером.", style='List Number')
    
    # Сохраняем документ
    filename = f"test_native_lists_{int(time.time())}.docx"
    doc.save(filename)
    return filename

def test_native_list_detection():
    """Тестирование определения встроенных элементов списка"""
    print("=== Тест определения встроенных элементов списка ===")
    test_file = create_test_doc_with_native_lists()
    
    doc = Document(test_file)
    
    print("Анализ документа на наличие встроенных элементов списка:")
    list_items = []
    
    for i, para in enumerate(doc.paragraphs):
        if is_list_item(para):
            list_items.append(i)
            print(f"[{i}] Распознан элемент списка: '{para.text}'")
    
    print(f"Всего найдено элементов списка: {len(list_items)}")
    
    return doc, list_items

def test_native_list_format_checking():
    """Тестирование проверки форматирования встроенных элементов списка"""
    print("\n=== Тест проверки форматирования встроенных элементов списка ===")
    doc, list_items = test_native_list_detection()
    
    comments_list = []
    
    # Создаем тестовый список параграфов с элементами списка
    test_paras = list(doc.paragraphs)
    
    # Отладочные данные о распознавании типов списков
    print("\n--- Диагностика определения типов списков ---")
    for para_idx in list_items:
        para = doc.paragraphs[para_idx]
        visible_text = para.text.lstrip()
        style_name = para.style.name.lower() if para.style and para.style.name else "нет стиля"
        
        # Базовые проверки для отображения диагностики
        bullet_match = re.search(r'^[-•*]', visible_text)
        number_match = re.search(r'^\d+[.)]\s', visible_text)
        letter_match = re.search(r'^[а-яa-z][.)]\s', visible_text)
        
        has_bullet_marker = bullet_match is not None
        has_number_marker = number_match is not None
        has_letter_marker = letter_match is not None
        
        print(f"\nПараграф #{para_idx}: '{visible_text[:30]}{'...' if len(visible_text) > 30 else ''}'")
        print(f"Стиль параграфа: {style_name}")
        print(f"Видимые маркеры: Маркер '-': {'✓' if has_bullet_marker else '✗'}, Цифра: {'✓' if has_number_marker else '✗'}, Буква: {'✓' if has_letter_marker else '✗'}")
    
    # Проверяем форматирование каждого найденного элемента списка
    for para_idx in list_items:
        check_list_item_format(doc.paragraphs[para_idx], para_idx, comments_list, "Test", test_paras, para_idx)
    
    # Выводим все комментарии
    print(f"\nНайдено {len(comments_list)} замечаний:")
    for idx, comment, author in comments_list:
        print(f"[{idx}] {comment}")
    
    return comments_list

if __name__ == "__main__":
    print("Создание тестового документа со встроенными списками...")
    test_file = create_test_doc_with_native_lists()
    print(f"Тестовый документ создан: {test_file}")
    
    test_native_list_detection()
    test_native_list_format_checking() 