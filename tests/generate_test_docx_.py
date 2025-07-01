from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE # Хотя WD_STYLE_TYPE не используется явно после удаления 'CommentText'

def create_test_document(filename="test_normcontrol_document.docx"):
    doc = Document()

    # --- 0. Настройка полей (для всего документа) ---
    section = doc.sections[0]
    # Неправильные поля (для теста)
    section.left_margin = Mm(25)
    section.right_margin = Mm(20)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    doc.add_paragraph("Этот абзац на странице с НЕПРАВИЛЬНЫМИ полями (25,20,25,25).")
    doc.add_paragraph() # Пустая строка для отделения

    # --- Титульный лист (должен игнорироваться большинством проверок) ---
    doc.add_paragraph("ТИТУЛЬНЫЙ ЛИСТ", style='Title') # Стиль 'Title' обычно существует
    para_title = doc.add_paragraph("Работа студента Иванова И.И.")
    para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if para_title.runs: # Проверка, что раны существуют
        run_title = para_title.runs[0]
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(16)
        run_title.font.bold = True
    doc.add_paragraph("\n" * 3) # Несколько пустых строк

    # --- Основной заголовок "ВВЕДЕНИЕ" (начало обработки) ---
    para = doc.add_paragraph("ВВЕДЕНИЕ")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    # --- Основной текст ---
    doc.add_paragraph("Правильный основной текст:")
    para = doc.add_paragraph(
        "Это пример основного текста. Он написан шрифтом Times New Roman, 14 кеглем, "
        "с полуторным межстрочным интервалом и абзацным отступом 1.25 см. "
        "Выравнивание текста должно быть по ширине. Цвет шрифта черный."
    )
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    doc.add_paragraph("Неправильный основной текст (шрифт, размер, интервал, отступ, выравнивание, цвет):")
    para = doc.add_paragraph(
        "Этот текст с ошибками. Шрифт Arial, 12 кегль, одинарный интервал, "
        "без абзацного отступа. Выравнивание по левому краю. Цвет шрифта синий."
    )
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.first_line_indent = Cm(0.5) 
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in para.runs:
        run.font.name = 'Arial' 
        run.font.size = Pt(12)  
        run.font.color.rgb = RGBColor(0,0,255)
    doc.add_paragraph()

    # --- Заголовок раздела ---
    doc.add_paragraph("--- Заголовки разделов ---")
    para = doc.add_paragraph() 
    para.paragraph_format.page_break_before = True
    para = doc.add_paragraph("1. Правильный заголовок раздела") 
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()
    
    doc.add_paragraph() 
    para = doc.add_paragraph("2. Неправильный заголовок раздела.") 
    para.paragraph_format.first_line_indent = Cm(1.0) 
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    for run in para.runs:
        run.font.name = 'Calibri' 
        run.font.size = Pt(16)
        run.bold = False 
        run.font.color.rgb = RGBColor(0,128,0)
    doc.add_paragraph()

    # --- Заголовок подраздела ---
    doc.add_paragraph("--- Заголовки подразделов ---")
    para = doc.add_paragraph("1.1. Правильный заголовок подраздела") 
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    para = doc.add_paragraph("1.1.1. Неправильный заголовок подраздела.") 
    para.paragraph_format.first_line_indent = Cm(0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = False 
        run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    # --- Подписи к рисункам и таблицам ---
    doc.add_paragraph("--- Рисунки и таблицы ---")
    para = doc.add_paragraph("Рисунок 1 – Правильное описание рисунка")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    para = doc.add_paragraph("Рисунок 2 – Неправильное описание.")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12) 
        run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    para = doc.add_paragraph("Таблица 1 – Правильное название таблицы")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # para.paragraph_format.first_line_indent = None # Установка в None может вызвать проблемы, лучше 0
    para.paragraph_format.first_line_indent = Cm(0)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0,0,0)
    
    table = doc.add_table(rows=1, cols=2)
    table.cell(0,0).text = "Текст в таблице 1"
    table.cell(0,1).text = "Текст в таблице 2 (Arial, 10pt)"
    if table.cell(0,1).paragraphs:
        for run in table.cell(0,1).paragraphs[0].runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
    doc.add_paragraph()


    para = doc.add_paragraph("Таблица 2 – Неправильное название таблицы.")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    para.paragraph_format.first_line_indent = Cm(1.25) 
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    # --- Списки ---
    doc.add_paragraph("--- Списки ---")
    para = doc.add_paragraph("- Первый элемент правильного списка")
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0,0,0)
    
    para = doc.add_paragraph("а) Второй элемент правильного списка (буква)")
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0,0,0)
    
    para = doc.add_paragraph("1) Третий элемент правильного списка (цифра)")
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    para = doc.add_paragraph("- Первый элемент неправильного списка (без отступа, другой шрифт)")
    para.paragraph_format.first_line_indent = Cm(0) 
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in para.runs: run.font.name = 'Calibri'; run.font.size = Pt(12); run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    # --- Список литературы ---
    para = doc.add_paragraph() 
    para.paragraph_format.page_break_before = True
    para = doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ") 
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.bold = True; run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    para = doc.add_paragraph("1. Иванов И.И. Моя книга. – М.: Наука, 2023. – 300 с.")
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0,0,0)
    
    para = doc.add_paragraph("Петров П.П. Другая статья // Журнал. 2022. №1. С. 5-10.") # Без номера
    para.paragraph_format.first_line_indent = Cm(0) 
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs: run.font.name = 'Arial'; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    # --- Ссылки в тексте ---
    doc.add_paragraph("--- Ссылки в тексте ---")
    doc.add_paragraph("Текст с правильной ссылкой [1]. И еще одна [2, с. 55].")
    doc.add_paragraph("Текст с неправильной ссылкой [ 3 ] и [4,с.10] и [5, с. 12-13].")
    doc.add_paragraph()

    # --- Сноски (имитация) ---
    doc.add_paragraph("--- Сноски (имитация) ---")
    para = doc.add_paragraph("Это текст, который мог бы содержать ссылку на сноску.")
    
    doc.add_paragraph("Текст имитации сноски 1 (Times New Roman, 10pt, одинарный)") 
    para_fn_good = doc.paragraphs[-1]
    for run in para_fn_good.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0,0,0)
    para_fn_good.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    doc.add_paragraph("Текст имитации сноски 2 (Arial, 14pt, полуторный)")
    para_fn_bad = doc.paragraphs[-1]
    for run in para_fn_bad.runs:
        run.font.name = 'Arial' 
        run.font.size = Pt(14)  
        run.font.color.rgb = RGBColor(0,0,255) # Синий
    para_fn_bad.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    doc.add_paragraph("Примечание: Реальные сноски для теста лучше добавить вручную в Word.")
    doc.add_paragraph()

    # --- Приложение ---
    para = doc.add_paragraph()
    para.paragraph_format.page_break_before = True
    para = doc.add_paragraph("ПРИЛОЖЕНИЕ А") 
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.bold = True; run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()
    
    para = doc.add_paragraph("Содержимое приложения А.")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    para = doc.add_paragraph("приложение б") # Неправильно
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs: run.font.name = 'Calibri'; run.font.size = Pt(12); run.bold = False; run.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

    # --- Конец документа ---
    doc.add_paragraph("КОНЕЦ ТЕСТОВОГО ДОКУМЕНТА")
    
    doc.save(filename)
    print(f"Тестовый документ '{filename}' успешно создан.")

if __name__ == '__main__':
    create_test_document()