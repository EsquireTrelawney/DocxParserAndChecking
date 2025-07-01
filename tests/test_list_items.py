#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from formatting_checker import is_list_item, check_list_item_format
import time

def create_test_doc_with_lists():
    """Создает тестовый документ с различными видами списков"""
    doc = Document()
    
    # Добавляем заголовок ВВЕДЕНИЕ для активации проверок
    heading = doc.add_paragraph("ВВЕДЕНИЕ")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.runs[0].bold = True
    
    # Пустая строка
    doc.add_paragraph()
    
    # Добавляем обычный текст
    doc.add_paragraph("Это обычный текст параграфа, который не является элементом списка.")
    
    # Заголовок для раздела со списками
    subsection = doc.add_paragraph("Тестовые списки")
    subsection.runs[0].bold = True
    
    # Правильный список с дефисами
    doc.add_paragraph("Правильные элементы списка:")
    
    # Правильные элементы списка
    p1 = doc.add_paragraph("- Элемент списка с дефисом")
    p1.paragraph_format.first_line_indent = Pt(14)  # Примерно 1.25 см
    
    p2 = doc.add_paragraph("а) Элемент списка с буквой и скобкой")
    p2.paragraph_format.first_line_indent = Pt(14)
    
    p3 = doc.add_paragraph("1) Элемент списка с цифрой и скобкой")
    p3.paragraph_format.first_line_indent = Pt(14)
    
    p4 = doc.add_paragraph("• Элемент списка с маркером")
    p4.paragraph_format.first_line_indent = Pt(14)
    
    p5 = doc.add_paragraph("1. Элемент списка с цифрой и точкой")
    p5.paragraph_format.first_line_indent = Pt(14)
    
    p6 = doc.add_paragraph("a. Элемент списка с буквой и точкой")
    p6.paragraph_format.first_line_indent = Pt(14)
    
    # Пустая строка
    doc.add_paragraph()
    
    doc.add_paragraph("Неправильные элементы списка:")
    
    # Неправильные элементы списка
    p7 = doc.add_paragraph("- Элемент списка с точкой в конце.")
    p7.paragraph_format.first_line_indent = Pt(14)
    
    p8 = doc.add_paragraph("а) Элемент списка с неправильным шрифтом")
    p8.paragraph_format.first_line_indent = Pt(14)
    p8.runs[0].font.name = "Arial"
    
    p9 = doc.add_paragraph("1) Элемент списка с неправильным размером шрифта")
    p9.paragraph_format.first_line_indent = Pt(14)
    p9.runs[0].font.size = Pt(12)
    
    p10 = doc.add_paragraph("* Элемент списка с неправильным маркером")
    p10.paragraph_format.first_line_indent = Pt(14)
    
    p11 = doc.add_paragraph("Элемент без маркера, но с отступом")
    p11.paragraph_format.first_line_indent = Pt(14)
    
    p12 = doc.add_paragraph("- Элемент списка с красным цветом")
    p12.paragraph_format.first_line_indent = Pt(14)
    p12.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    # Добавляем элементы с неправильными маркерами, но с правильным отступом
    p13 = doc.add_paragraph("1. Неправильный элемент списка с цифрой и точкой")
    p13.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    p14 = doc.add_paragraph("a. Неправильный элемент списка с буквой и точкой")
    p14.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    # Добавляем группу элементов списка для проверки правил точки с запятой
    doc.add_paragraph()
    doc.add_paragraph("Группа элементов списка с правильными и неправильными окончаниями:")
    
    # Правильное форматирование (точка с запятой у всех кроме последнего)
    p15 = doc.add_paragraph("- Первый элемент правильного списка;")
    p15.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    p16 = doc.add_paragraph("- Второй элемент правильного списка;")
    p16.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    p17 = doc.add_paragraph("- Третий элемент правильного списка.")
    p17.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    # Неправильное форматирование (точка вместо точки с запятой между элементами)
    doc.add_paragraph()
    p18 = doc.add_paragraph("- Первый элемент неправильного списка.")
    p18.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    p19 = doc.add_paragraph("- Второй элемент неправильного списка.")
    p19.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    p20 = doc.add_paragraph("- Третий элемент неправильного списка.")
    p20.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    # Неправильное форматирование (нет точки с запятой в середине и нет точки в конце)
    doc.add_paragraph()
    p21 = doc.add_paragraph("- Первый элемент списка без точки с запятой")
    p21.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    p22 = doc.add_paragraph("- Последний элемент списка без точки")
    p22.paragraph_format.first_line_indent = Pt(14)  # 1.25 см
    
    # Сохраняем документ
    filename = f"test_lists_{int(time.time())}.docx"
    doc.save(filename)
    return filename

def test_list_detection():
    """Тестирование определения элементов списка"""
    print("=== Тест определения элементов списка ===")
    test_file = create_test_doc_with_lists()
    
    doc = Document(test_file)
    
    print("Анализ документа на наличие элементов списка:")
    list_items = []
    
    for i, para in enumerate(doc.paragraphs):
        if is_list_item(para):
            list_items.append(i)
            print(f"[{i}] Распознан элемент списка: '{para.text}'")
    
    print(f"Всего найдено элементов списка: {len(list_items)}")
    
    return doc, list_items

def test_list_format_checking():
    """Тестирование проверки форматирования элементов списка"""
    print("\n=== Тест проверки форматирования элементов списка ===")
    doc, list_items = test_list_detection()
    
    comments_list = []
    
    # Создаем тестовый список параграфов с элементами списка
    test_paras = []
    for i in range(len(doc.paragraphs)):
        test_paras.append(doc.paragraphs[i])
    
    # Проверяем форматирование каждого найденного элемента списка
    for para_idx in list_items:
        check_list_item_format(doc.paragraphs[para_idx], para_idx, comments_list, "Test", test_paras, para_idx)
    
    # Выводим все комментарии
    print(f"Найдено {len(comments_list)} замечаний:")
    for idx, comment, author in comments_list:
        print(f"[{idx}] {comment}")
    
    return comments_list

if __name__ == "__main__":
    print("Создание тестового документа со списками...")
    test_file = create_test_doc_with_lists()
    print(f"Тестовый документ создан: {test_file}")
    
    test_list_detection()
    test_list_format_checking() 