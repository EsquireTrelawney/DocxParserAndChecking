#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
import sys
import re
import traceback
from docx import Document
from formatting_checker import (
    is_list_item, check_list_item_format, is_bibliography_item, 
    is_bibliography_heading, check_bibliography_item_format, 
    check_gost_bibliography_compliance, check_bibliography_numbering,
    is_main_heading
)

def analyze_paragraph(doc, para_idx, force_list_check=False):
    """
    Анализирует конкретный параграф из документа Word.
    
    Args:
        doc: объект Document из python-docx
        para_idx: индекс параграфа для анализа
        force_list_check: проверять форматирование как для элемента списка, даже если параграф не распознан как список
    """
    if para_idx < 0 or para_idx >= len(doc.paragraphs):
        print(f"Ошибка: Индекс параграфа {para_idx} вне диапазона (0-{len(doc.paragraphs)-1})")
        return
    
    try:
        para = doc.paragraphs[para_idx]
        text = para.text.strip()
        visible_text = ''.join(ch for ch in text if ch.isprintable())
        
        # Проверка, является ли параграф элементом списка
        is_list = is_list_item(para)
        
        # Базовая информация о параграфе
        print(f"\n--- Параграф #{para_idx}: '{text[:50]}{'...' if len(text) > 50 else ''}'")
        print(f"Распознан как элемент списка: {'Да' if is_list else 'Нет'}")
        
        # Проверка стиля параграфа
        style_name = "не определено"
        if hasattr(para, 'style') and para.style:
            if hasattr(para.style, 'name') and para.style.name:
                style_name = para.style.name
        print(f"Стиль параграфа: {style_name}")
        
        # Проверка атрибутов форматирования
        if hasattr(para, 'paragraph_format') and para.paragraph_format:
            pf = para.paragraph_format
            
            # Отступ первой строки
            first_indent = "не установлен"
            if hasattr(pf, 'first_line_indent') and pf.first_line_indent:
                first_indent = f"{pf.first_line_indent.cm:.2f} см"
            print(f"Отступ первой строки: {first_indent}")
            
            # Левый отступ
            left_indent = "не установлен"
            if hasattr(pf, 'left_indent') and pf.left_indent:
                left_indent = f"{pf.left_indent.cm:.2f} см"
            print(f"Левый отступ: {left_indent}")
            
            # Информация о нумерации
            if hasattr(pf, 'numbering') and pf.numbering:
                num_info = []
                if hasattr(pf.numbering, 'level') and pf.numbering.level is not None:
                    num_info.append(f"уровень={pf.numbering.level}")
                if hasattr(pf.numbering, 'num_id') and pf.numbering.num_id is not None:
                    num_info.append(f"num_id={pf.numbering.num_id}")
                print(f"Атрибуты нумерации: {', '.join(num_info)}")
            else:
                print("Атрибуты нумерации: отсутствуют")
        
        # Проверка начала текста (маркеры списка)
        print(f"Начало текста (без пробелов): '{visible_text.lstrip()[:10]}'")
        
        # Проверка на соответствие шаблонам маркеров списка
        patterns = {
            "Дефис (-)": r"^-\s+",
            "Буква со скобкой (а))": r"^\w\)\s+",
            "Цифра со скобкой (1))": r"^\d+\)\s+",
            "Цифра с точкой (1.)": r"^\d+\.\s+",
            "Буква с точкой (а.)": r"^[а-яА-Яa-zA-Z]\.\s+",
            "Звездочка (*)": r"^[*]\s+",
            "Специальный маркер (•)": r"^[•●○◦■□▪▫]\s+"
        }
        
        print("Проверка на соответствие шаблонам маркеров:")
        for name, pattern in patterns.items():
            match = re.match(pattern, visible_text.lstrip())
            print(f"  - {name}: {'✓' if match else '✗'}")
        
        # Дополнительная проверка для элементов с нумерацией
        if hasattr(para, 'paragraph_format') and para.paragraph_format and \
           hasattr(para.paragraph_format, 'numbering') and para.paragraph_format.numbering:
            if hasattr(para.paragraph_format.numbering, 'num_id') and para.paragraph_format.numbering.num_id:
                # Получаем первые 100 символов текста для анализа
                text_prefix = visible_text.lstrip()[:100]
                print(f"\nДетальный анализ текста с нумерацией:")
                print(f"Первые 100 символов: '{text_prefix}'")
                
                # Проверка на каждый символ (помогает найти невидимые или специальные символы)
                print("Коды символов:")
                char_codes = [(i, c, ord(c)) for i, c in enumerate(text_prefix[:20])]
                for i, c, code in char_codes:
                    print(f"  Позиция {i}: '{c}' (код {code}, hex: {hex(code)})")
        
        # Проверка форматирования как для элемента списка (даже если параграф не распознан как список)
        if is_list or force_list_check:
            print("\nПроверка форматирования элемента списка:")
            try:
                comments = []
                check_list_item_format(para, para_idx, comments, "Test", doc.paragraphs, para_idx)
                if comments:
                    for _, comment, _ in comments:
                        print(f"  - {comment}")
                else:
                    print("  - Проблем с форматированием не обнаружено")
                    
                # Добавляем подробный вывод типа списка и контекста    
                print("\nОтладочная информация для определения типа списка:")
                # Добавляем информацию о соседних параграфах
                if para_idx > 0 and para_idx < len(doc.paragraphs) - 1:
                    prev_para = doc.paragraphs[para_idx - 1]
                    next_para = doc.paragraphs[para_idx + 1]
                    
                    prev_style = "не определен"
                    if hasattr(prev_para, 'style') and prev_para.style and hasattr(prev_para.style, 'name'):
                        prev_style = prev_para.style.name
                    
                    next_style = "не определен"
                    if hasattr(next_para, 'style') and next_para.style and hasattr(next_para.style, 'name'):
                        next_style = next_para.style.name
                        
                    print(f"  * Предыдущий параграф: стиль '{prev_style}', текст: '{prev_para.text[:30]}...'")
                    print(f"  * Следующий параграф: стиль '{next_style}', текст: '{next_para.text[:30]}...'")
                    
                    # Проверка на наличие маркеров в соседних параграфах
                    prev_has_bullet = re.search(r'^[-•*]', prev_para.text.lstrip()) is not None
                    prev_has_number = re.search(r'^\d+[.)]\s', prev_para.text.lstrip()) is not None
                    prev_has_letter = re.search(r'^[а-яa-z][.)]\s', prev_para.text.lstrip()) is not None
                    
                    next_has_bullet = re.search(r'^[-•*]', next_para.text.lstrip()) is not None
                    next_has_number = re.search(r'^\d+[.)]\s', next_para.text.lstrip()) is not None
                    next_has_letter = re.search(r'^[а-яa-z][.)]\s', next_para.text.lstrip()) is not None
                    
                    print(f"  * Маркеры в предыдущем: дефис={prev_has_bullet}, цифра={prev_has_number}, буква={prev_has_letter}")
                    print(f"  * Маркеры в следующем: дефис={next_has_bullet}, цифра={next_has_number}, буква={next_has_letter}")
                    
                    # Определяем тип списка по контексту
                    if prev_has_number or prev_has_letter or next_has_number or next_has_letter:
                        print("  * Предположительно нумерованный список (по контексту)")
                    elif prev_has_bullet or next_has_bullet:
                        print("  * Предположительно маркированный список (по контексту)")
            except Exception as e:
                print(f"  - Ошибка при проверке форматирования: {str(e)}")
                # Детальный анализ ошибки
                print("\nДетальная информация об ошибке:")
                print(traceback.format_exc())
                
                # Дополнительная проверка на проблемы с para.runs
                if 'para.runs' in str(e) or 'NoneType' in str(e):
                    print("\nДополнительная проверка атрибутов параграфа:")
                    print(f"Наличие параграфа: {'Да' if para else 'Нет'}")
                    if para:
                        print(f"Наличие runs: {'Да' if hasattr(para, 'runs') else 'Нет'}")
                        if hasattr(para, 'runs'):
                            print(f"Количество runs: {len(para.runs)}")
                            print("Информация о каждом run:")
                            for i, run in enumerate(para.runs):
                                print(f"  Run #{i}:")
                                print(f"    Текст: '{run.text if hasattr(run, 'text') else 'Нет текста'}'")
                                if hasattr(run, 'font'):
                                    print(f"    Font: {'Есть' if run.font else 'Отсутствует'}")
                                    if run.font and hasattr(run.font, 'name'):
                                        print(f"    Font name: {run.font.name}")
        
        print("-" * 60)
    except Exception as e:
        print(f"Ошибка при анализе параграфа #{para_idx}: {str(e)}")
        print(traceback.format_exc())

def test_native_and_real_lists(docx_path, verbose=False, para_indices=None):
    """Тестирование обработки и проверки элементов списков в документе Word"""
    try:
        doc = Document(docx_path)
        print(f"Документ успешно загружен, всего параграфов: {len(doc.paragraphs)}")
        
        # Анализ конкретных параграфов, если они указаны
        if para_indices:
            print(f"Анализ конкретных параграфов: {para_indices}")
            for para_idx in para_indices:
                analyze_paragraph(doc, para_idx)
        # Или анализ всех параграфов с подробным выводом информации
        elif verbose:
            print("\n=== Анализ всех элементов списков в документе ===")
            list_count = 0
            
            for i, para in enumerate(doc.paragraphs):
                if is_list_item(para):
                    list_count += 1
                    analyze_paragraph(doc, i)
            
            print(f"\nВсего найдено элементов списка: {list_count}")
        # Или просто вывод информации о найденных элементах списка
        else:
            list_items = []
            for i, para in enumerate(doc.paragraphs):
                if is_list_item(para):
                    list_items.append(i)
            
            print(f"\nНайдено {len(list_items)} элементов списка: {list_items}")
                    
    except Exception as e:
        print(f"Ошибка при анализе документа: {str(e)}")
        print(traceback.format_exc())

def test_bibliography_items(doc_path, verbose=False, para_indices=None):
    """Тестирование обработки и проверки элементов библиографии"""
    print(f"\n=== Тест проверки форматирования элементов библиографии ===")
    
    # Загрузить документ
    doc = Document(doc_path)
    print(f"Документ успешно загружен, всего параграфов: {len(doc.paragraphs)}")
    
    # Находим секцию библиографии
    bibliography_sections = []
    bibliography_start_idx = -1
    bibliography_items = []
    
    # Перебираем все параграфы, чтобы найти все заголовки библиографии
    for i, para in enumerate(doc.paragraphs):
        if is_bibliography_heading(para):
            bibliography_sections.append((i, para.text))
            print(f"\nНайден заголовок библиографии: '{para.text}' (индекс: {i})")
    
    # Если не нашли ни одного заголовка библиографии
    if not bibliography_sections:
        print("Не найдено ни одного заголовка библиографии")
        return
    
    # Берем последний заголовок библиографии как основной
    bibliography_start_idx = bibliography_sections[-1][0]
    print(f"Используем в качестве основного заголовка библиографии: '{bibliography_sections[-1][1]}' (индекс: {bibliography_start_idx})")
    
    # Находим элементы библиографии после этого заголовка
    in_bibliography_section = False
    
    for i, para in enumerate(doc.paragraphs):
        # Включаем режим библиографии, когда встречаем последний заголовок
        if i == bibliography_start_idx:
            in_bibliography_section = True
            continue
            
        # Если встретился другой структурный заголовок после заголовка библиографии
        if in_bibliography_section and is_main_heading(para) and not is_bibliography_heading(para):
            in_bibliography_section = False
            continue
            
        # Если мы в разделе библиографии и это элемент библиографии
        if in_bibliography_section and is_bibliography_item(para, True):
            bibliography_items.append(i)
    
    # Выводим информацию о найденных элементах
    if not bibliography_items:
        print("Не найдено ни одного элемента библиографии")
        return
        
    print(f"Найдено {len(bibliography_items)} элементов библиографии")
    
    # Если указаны конкретные параграфы для проверки, используем их
    if para_indices:
        items_to_check = para_indices
    else:
        items_to_check = bibliography_items
    
    # Проверяем каждый элемент
    for para_idx in items_to_check:
        if para_idx >= len(doc.paragraphs):
            print(f"Ошибка: параграф #{para_idx} не существует")
            continue
            
        para = doc.paragraphs[para_idx]
        text = para.text.strip()
        
        print(f"\n--- Параграф #{para_idx}: '{text[:40]}...'")
        
        # Проверяем, находится ли параграф в разделе библиографии
        is_in_bibliography = para_idx > bibliography_start_idx
        
        # Для параграфов до начала библиографии, явно указываем, что они вне библиографии
        if not is_in_bibliography:
            is_biblio = False
            print(f"Распознан как элемент библиографии: Нет (находится до раздела библиографии)")
        else:
            # Проверяем, распознан ли как элемент библиографии
            is_biblio = is_bibliography_item(para, True)
            print(f"Распознан как элемент библиографии: {'Да' if is_biblio else 'Нет'}")
        
        # Если это не элемент библиографии, продолжаем со следующим параграфом
        if not is_biblio:
            continue
        
        # Анализ стиля параграфа
        style_name = para.style.name if hasattr(para, 'style') and para.style and hasattr(para.style, 'name') else "не определен"
        print(f"Стиль параграфа: {style_name}")
        
        # Анализ отступов
        first_indent = para.paragraph_format.first_line_indent.cm if hasattr(para, 'paragraph_format') and para.paragraph_format and hasattr(para.paragraph_format, 'first_line_indent') and para.paragraph_format.first_line_indent else "не установлен"
        print(f"Отступ первой строки: {first_indent}")
        
        # Проверка соответствия формату "N. Текст"
        has_number = re.match(r"^\d+\.\s+", text) is not None
        print(f"Имеет формат 'N. Текст': {'Да' if has_number else 'Нет'}")
        
        if has_number:
            number = re.match(r"^(\d+)\.", text).group(1)
            print(f"Номер элемента библиографии: {number}")
        
        # Проводим проверку форматирования
        comments = []
        check_bibliography_item_format(para, para_idx, comments, "Test")
        
        print("Проверка форматирования элемента библиографии:")
        if comments:
            for _, comment, _ in comments:
                print(f"  - {comment}")
        else:
            print("  - Проблем с форматированием не обнаружено")
            
        # Проводим проверку соответствия ГОСТ
        comments = []
        check_gost_bibliography_compliance(text, para_idx, comments, "Test")
        
        print("Проверка соответствия ГОСТ Р 7.0.100-2018:")
        if comments:
            for _, comment, _ in comments:
                print(f"  - {comment}")
        else:
            print("  - Соответствует требованиям ГОСТ")
        
        print("------------------------------------------------------------")
    
    # Проверка последовательности нумерации
    if verbose and bibliography_start_idx >= 0:
        print("\n--- Проверка последовательности нумерации библиографических записей ---")
        comments = []
        check_bibliography_numbering(doc.paragraphs, bibliography_start_idx, comments, "Test")
        
        if comments:
            for idx, comment, _ in comments:
                print(f"  - Параграф #{idx}: {comment}")
        else:
            print("  - Проблем с нумерацией не обнаружено")

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Анализ списков в документе Word')
    parser.add_argument('docx_path', type=str, help='Путь к документу Word')
    parser.add_argument('-v', '--verbose', action='store_true', help='Подробный вывод')
    parser.add_argument('-p', '--paragraph', type=int, action='append', dest='paragraphs',
                        help='Индекс параграфа для анализа (можно указать несколько раз)')
    parser.add_argument('-b', '--bibliography', action='store_true', 
                        help='Проверить элементы библиографического списка')
    
    args = parser.parse_args()
    
    try:
        print(f"Анализ списков в документе: {args.docx_path}")
        
        if args.bibliography:
            test_bibliography_items(args.docx_path, args.verbose, args.paragraphs)
        else:
            test_native_and_real_lists(args.docx_path, args.verbose, args.paragraphs)
            
    except Exception as e:
        print(f"Ошибка при выполнении: {e}")
        import traceback
        print(traceback.format_exc()) 