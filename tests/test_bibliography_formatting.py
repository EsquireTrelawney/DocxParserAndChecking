#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Тестирование функций проверки библиографии
"""

import os
import sys
from docx import Document
from formatting_checker import (
    is_bibliography_item, is_bibliography_heading, check_bibliography_item_format,
    check_gost_bibliography_compliance, check_bibliography_numbering,
    check_document_formatting_final
)

def test_bibliography_detection(doc_path):
    """Я написал функцию для проверки обнаружения библиографических записей"""
    print(f"\n=== Тестирование обнаружения элементов библиографии в {os.path.basename(doc_path)} ===")
    
    doc = Document(doc_path)
    
    # Находим заголовок библиографии
    bibliography_headings = []
    for i, para in enumerate(doc.paragraphs):
        if is_bibliography_heading(para):
            bibliography_headings.append((i, para.text))
    
    print(f"Найдено заголовков библиографии: {len(bibliography_headings)}")
    for idx, text in bibliography_headings:
        print(f"  - Параграф #{idx}: '{text}'")
    
    # Находим элементы библиографии
    bibliography_items = []
    in_bibliography_section = False
    
    for i, para in enumerate(doc.paragraphs):
        if is_bibliography_heading(para):
            in_bibliography_section = True
            continue
            
        if in_bibliography_section and is_bibliography_item(para, True):
            bibliography_items.append((i, para.text[:40] + "..."))
    
    print(f"\nНайдено элементов библиографии: {len(bibliography_items)}")
    for idx, text in bibliography_items:
        print(f"  - Параграф #{idx}: '{text}'")
            
def test_bibliography_checking(doc_path):
    """Тут я проверяю правильность форматирования библиографии"""
    print(f"\n=== Тестирование проверки форматирования элементов библиографии в {os.path.basename(doc_path)} ===")
    
    # Получаем все ошибки форматирования
    all_errors = check_document_formatting_final(doc_path)
    
    # Фильтруем только ошибки, относящиеся к библиографии
    bibliography_errors = [
        (i, comment) for i, comment, _ in all_errors 
        if "библиографич" in comment.lower() 
        or "список литературы" in comment.lower()
        or "список использованных источников" in comment.lower()
    ]
    
    print(f"Всего ошибок форматирования: {len(all_errors)}")
    print(f"Ошибок форматирования библиографии: {len(bibliography_errors)}")
    
    if bibliography_errors:
        print("\nОшибки форматирования библиографии:")
        for i, comment in bibliography_errors:
            print(f"  - Параграф #{i}: {comment}")
    else:
        print("\nОшибок форматирования библиографии не найдено")
        
def main():
    """Основная функция"""
    if len(sys.argv) < 2:
        print("Использование: python test_bibliography_formatting.py <путь_к_docx_файлу>")
        return
    
    doc_path = sys.argv[1]
    if not os.path.exists(doc_path):
        print(f"Ошибка: файл {doc_path} не найден")
        return
    
    test_bibliography_detection(doc_path)
    test_bibliography_checking(doc_path)

if __name__ == "__main__":
    main() 