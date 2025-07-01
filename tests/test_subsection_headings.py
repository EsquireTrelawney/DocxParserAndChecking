from docx import Document
from formatting_checker import is_subsection_heading, check_subsection_heading_format

# Создаем документ для тестирования
doc = Document()

# Создаем два параграфа для тестирования
correct = doc.add_paragraph("1.1 Правильный подзаголовок")
correct.runs[0].bold = True

incorrect = doc.add_paragraph("1.1. Неправильный подзаголовок с точкой")
incorrect.runs[0].bold = True

# Тест распознавания
print("==== Тест распознавания подзаголовков ====")
print(f"1.1 Правильный подзаголовок - распознан: {is_subsection_heading(correct)}")
print(f"1.1. Неправильный подзаголовок с точкой - распознан: {is_subsection_heading(incorrect)}")

# Тест проверки форматирования
print("\n==== Тест проверки форматирования подзаголовков ====")

# Проверка правильного подзаголовка
comments = []
check_subsection_heading_format(correct, 0, comments, "Test")
print("\nКомментарии для правильного подзаголовка (1.1 Название):")
if not comments:
    print("Нет замечаний (правильно)")
else:
    for idx, comment, author in comments:
        print(f"[{idx}] {comment}")

# Проверка неправильного подзаголовка (с точкой)
comments = []
check_subsection_heading_format(incorrect, 1, comments, "Test")
print("\nКомментарии для неправильного подзаголовка (1.1. Название):")
if not comments:
    print("Нет замечаний (неверно! должно быть замечание о точке)")
else:
    for idx, comment, author in comments:
        print(f"[{idx}] {comment}")
        
# Проверка наличия комментария о точке
has_period_warning = any("точки" in comment for idx, comment, author in comments)
print(f"\nОбнаружено предупреждение о точке после номера: {has_period_warning}") 