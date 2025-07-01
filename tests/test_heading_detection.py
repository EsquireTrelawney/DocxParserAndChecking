from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from formatting_checker import (
    is_main_heading, is_introduction_heading, is_bibliography_heading,
    is_appendix_heading, check_document_formatting_final
)


def create_test_doc():
    doc = Document()

    # Правильные заголовки
    doc.add_paragraph("ВВЕДЕНИЕ", style="Heading 1").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Основной текст введения")
    
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ", style="Heading 1").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Основной текст заключения")
    
    # Неправильные заголовки (в нижнем регистре, с точкой)
    doc.add_paragraph("заключение.", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph("Это другой абзац")
    
    doc.add_paragraph("приложение а", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph("Содержимое приложения")
    
    test_file = "test_headings.docx"
    doc.save(test_file)
    return test_file

def test_heading_detection():
    doc = Document()
    
    # Правильные заголовки
    para = doc.add_paragraph("ВВЕДЕНИЕ")
    assert is_main_heading(para) == True
    assert is_introduction_heading(para) == True
    
    para = doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    assert is_main_heading(para) == True
    
    # Заголовки с точкой
    para = doc.add_paragraph("ЗАКЛЮЧЕНИЕ.")
    assert is_main_heading(para) == True
    
    # Заголовки в нижнем регистре
    para = doc.add_paragraph("заключение")
    assert is_main_heading(para) == True
    
    # Смешанные варианты
    para = doc.add_paragraph("Заключение.")
    assert is_main_heading(para) == True
    
    # Приложения
    para = doc.add_paragraph("ПРИЛОЖЕНИЕ А")
    assert is_appendix_heading(para) == True
    
    para = doc.add_paragraph("Приложение А.")
    assert is_appendix_heading(para) == True

def test_document_check():
    test_file = create_test_doc()
    results = check_document_formatting_final(test_file)
    
    # Вывести все комментарии
    print(f"Найдено {len(results)} комментариев:")
    for idx, comment, author in results:
        print(f"[{idx}] {comment}")
    
    # Проверить, что были найдены комментарии к заголовку "заключение."
    found_case_comment = False
    found_dot_comment = False
    
    for idx, comment, author in results:
        # Проверяем разные форматы сообщения - и со словом "заключение" и просто "структурный заголовок"
        if ("заключение" in comment.lower() and "верхнем регистре" in comment.lower()) or \
           ('"заключение."' in comment.lower() and "верхнем регистре" in comment.lower()):
            found_case_comment = True
            print(f"НАЙДЕНО сообщение о регистре: {comment}")
        
        if ("заключение" in comment.lower() and "точкой" in comment.lower()) or \
           ('"заключение."' in comment.lower() and "точкой" in comment.lower()):
            found_dot_comment = True
            print(f"НАЙДЕНО сообщение о точке: {comment}")
    
    assert found_case_comment, "Не найден комментарий о регистре для неправильного заголовка 'заключение.'"
    assert found_dot_comment, "Не найден комментарий о точке для неправильного заголовка 'заключение.'"

if __name__ == "__main__":
    test_heading_detection()
    test_document_check()
    print("Все тесты успешно пройдены!") 